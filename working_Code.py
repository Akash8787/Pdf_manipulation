# import requests
# import json
# from pdf2docx import Converter
# from docx2pdf import convert
# from num2words import num2words
# from docx.shared import Pt
# from pathlib import Path
# import os
# import pythoncom
# import time
# import locale
# from docx import Document

# locale.setlocale(locale.LC_ALL, '')

# def Licence():
#     error = ""
#     message = ""
#     url = "http://aisagar.com/DeepBluePDFManipulation"
#     payload = json.dumps({
#         "token": "JUYHJ-LKJH12-IKJUYH-45OIUJ"
#     })
#     headers = {
#         'Content-Type': 'application/json'
#     }
#     response = requests.request("POST", url, headers=headers, data=payload)    
#     obj = json.loads(response.text)
#     if "Error" in response.text and "Message" in response.text:
#         error = obj["Error"]
#         message = obj["Message"]
#     else:
#         error = response.text    
#     return error, message

# def pdftodocx(pdf_file_path, doc_file_path):
#     obj = Converter(pdf_file_path)
#     obj.convert(doc_file_path)
#     obj.close()

# def font_change(doc_file_path):
#     def change_font_size(doc, target_text, new_font_size):
#         for table in doc.tables:
#             for row in table.rows:
#                 for cell in row.cells:
#                     if target_text in cell.text:
#                         for paragraph in cell.paragraphs:
#                             for run in paragraph.runs:
#                                 run.font.size = Pt(new_font_size)

#     doc = Document(doc_file_path)
#     target_text = 'Port of Discharge:'
#     new_font_size = 9
#     change_font_size(doc, target_text, new_font_size)
#     doc.save(doc_file_path)

# def decrease_number_by_percent(number):
#     decrease_percentage = 40
#     decrease_factor = 1 - (decrease_percentage / 100)
#     decreased_number = number * decrease_factor
#     return float(decreased_number)

# def remove_comma(num):
#     return num.replace(",", '')

# def usd_change(amount):
#     a1 = str(amount)
#     q1, q2 = a1.split('.')
#     words1 = num2words(int(q1), lang='en').capitalize()
#     words2 = num2words(int(q2), lang='en').capitalize()
#     words1 = words1.replace(" and", '')
#     words1 += ' dollars and '
#     words2 += ' cents only'
#     word = words1.upper() + words2.upper()
#     return word

# def formatAmount(amt):
#     amt = round(amt, 2)
#     formatted_amt = f"{amt:.2f}"
#     print(f"formatAmount input: {amt}, output: {formatted_amt}")
#     return formatted_amt

# def is_number(s):
#     """Return True if s can be converted to float."""
#     try:
#         float(s.replace(",", ""))
#         return True
#     except ValueError:
#         return False

# def decrease_number_by_percent(value):
#     """Reduce the value by 40% (retain 60%)."""
#     return value * 0.6

# def read_tables_from_docx(doc):
#     total_qty = 0.0
#     total_amount = 0.0
#     total_steel = 0.0  # Accumulate decreased steel values
#     frt_amount = 0.0
#     ins_amount = 0.0
#     frt_yes = False
#     ins_yes = False
#     fob_yes = False
#     data_table_found = False
#     total_row = None  # Initialize total_row outside the loop

#     for table in doc.tables:
#         for row_idx, row in enumerate(table.rows):
#             cells = row.cells
#             print("Cells:", [cell.text.strip() for cell in cells])  # Debug print

#             # Check if the table has enough columns for a data table
#             if len(cells) < 12:  # Minimum columns required (adjusted for multi-page case)
#                 print(f"Skipping row {row_idx} due to insufficient columns: {len(cells)} columns")
#                 continue

#             # Detect start of data table
#             if not data_table_found and cells[0].text.strip() == "S.N":
#                 data_table_found = True
#                 continue  # Skip header row

#             if data_table_found:
#                 # Identify and store the total row
#                 if "Total:" in cells[0].text:
#                     total_row = row
#                     continue

#                 try:
#                     # Qty (adjust index based on structure)
#                     qty_text = cells[9].text.strip().replace(",", "") if len(cells) > 9 else "0"
#                     qty = float(qty_text) if is_number(qty_text) else 0.0
#                     total_qty += qty
#                     print(f"Row {row_idx}: Qty = {qty}, Total Qty = {total_qty}")

#                     # Rate in USD (col 16, adjust if needed)
#                     rate_text = cells[16].text.strip().replace(",", "") if len(cells) > 16 else ""
#                     if is_number(rate_text):
#                         rate = float(rate_text)
#                         decreased_rate = round(decrease_number_by_percent(rate), 6)
#                         for p in cells[16].paragraphs:
#                             for run in p.runs:
#                                 run.text = f"{decreased_rate:.6f}"
#                         new_amount = round(qty * decreased_rate, 2)
#                         total_amount += new_amount
#                         for p in cells[17].paragraphs:
#                             for run in p.runs:
#                                 run.text = formatAmount(new_amount)
#                         print(f"Row {row_idx}: Rate = {rate} -> {decreased_rate}, New Amount = {new_amount}, Total Amount = {total_amount}")

#                     # Rate of Steel Per KG (col 12)
#                     rate_steel_text = cells[12].text.strip().replace(",", "") if len(cells) > 12 else ""
#                     if is_number(rate_steel_text):
#                         rate_steel = float(rate_steel_text)
#                         decreased_rate_steel = round(decrease_number_by_percent(rate_steel), 6)
#                         for p in cells[12].paragraphs:
#                             for run in p.runs:
#                                 run.text = f"{decreased_rate_steel:.6f}"

#                     # Total value of Steel content used (col 13)
#                     steel_text = cells[13].text.strip().replace(",", "") if len(cells) > 13 else ""
#                     if is_number(steel_text):
#                         steel_val = float(steel_text)
#                         decreased_steel = round(decrease_number_by_percent(steel_val), 2)
#                         total_steel += decreased_steel  # Accumulate decreased steel value
#                         for p in cells[13].paragraphs:
#                             for run in p.runs:
#                                 run.text = formatAmount(decreased_steel)
#                         print(f"Row {row_idx}: Steel Val = {steel_val} -> {decreased_steel}, Total Steel = {total_steel}")

#                 except IndexError as e:
#                     print(f"Error processing row {row_idx}: {e}")
#                     continue

#     # After processing all rows and tables, update the total amount row
#     if total_row is not None:
#         # Dynamically find the indices of numeric values in the total row
#         numeric_values = [(i, float(cell.text.strip().replace(",", ""))) for i, cell in enumerate(total_row.cells) if is_number(cell.text.strip())]
#         if len(numeric_values) >= 2:
#             amount_idx = numeric_values[-1][0]  # Last numeric value (Grand Total in USD)
#             steel_idx = numeric_values[-2][0]   # Second-to-last numeric value (Total value of Steel content used)
#         else:
#             amount_idx = 17  # Fallback to default if not enough numeric values
#             steel_idx = 13

#         # Debug the indices and current values
#         print(f"Updating steel_idx = {steel_idx}, amount_idx = {amount_idx}")
#         print(f"Current Steel Value: {total_row.cells[steel_idx].text.strip()}")
#         print(f"Current Amount Value: {total_row.cells[amount_idx].text.strip()}")

#         # Update Grand Total in USD
#         has_paragraphs = False
#         for p in total_row.cells[amount_idx].paragraphs:
#             has_paragraphs = True
#             for run in p.runs:
#                 run.text = formatAmount(round(total_amount, 2))
#                 print(f"Setting Grand Total in USD (col {amount_idx}) to {formatAmount(round(total_amount, 2))}")
#         if not has_paragraphs:
#             print(f"No paragraphs found in Grand Total cell {amount_idx}")
#             # Fallback to direct text update if no paragraphs
#             total_row.cells[amount_idx].text = formatAmount(round(total_amount, 2))

#         # Update Total value of Steel content used
#         has_paragraphs = False
#         formatted_total_steel = formatAmount(round(total_steel, 2))
#         original_total_steel = total_row.cells[steel_idx].text.strip()
#         for p in total_row.cells[steel_idx].paragraphs:
#             has_paragraphs = True
#             for run in p.runs:
#                 run.text = run.text.replace(original_total_steel, formatted_total_steel)
#                 print(f"Setting Total value of Steel content used (col {steel_idx}) to {formatted_total_steel}")
#         if not has_paragraphs:
#             print(f"No paragraphs found in Steel Total cell {steel_idx}")
#             # Fallback to direct text update if no paragraphs
#             total_row.cells[steel_idx].text = formatted_total_steel

#         # Print the cell text after updates
#         print("Updated Total Row Cells:")
#         print(f"Total value of Steel content used (col {steel_idx}): {total_row.cells[steel_idx].text.strip()}")
#         print(f"Grand Total in USD (col {amount_idx}): {total_row.cells[amount_idx].text.strip()}")
 
#     # Update other sections (unchanged)
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 if 'Amount in words: USD' in cell.text:
#                     amount_words = cell.text.replace('Amount in words: USD ', '')
#                     for paragraph in cell.paragraphs:
#                         for run in paragraph.runs:
#                             run.text = run.text.replace(amount_words, usd_change(total_amount))

#                 if not frt_yes:    
#                     if 'FRT AMT IN USD -' in cell.text:
#                         x = cell.text.split('-')
#                         space = sum(1 for xx in x[-1] if xx == ' ') if len(x) > 1 else 0
#                         x = x[-1].lstrip()
#                         FRT = remove_comma(x)
#                         frt_amount = float(FRT)
#                         frt_amount = decrease_number_by_percent(frt_amount)
#                         total_amount -= frt_amount
#                         for paragraph in cell.paragraphs:
#                             for run in paragraph.runs:
#                                 run.text = run.text.replace(cell.text, f'FRT AMT IN USD -{(space * " ")+formatAmount(frt_amount)}')
#                         frt_yes = True

#                 if not ins_yes:
#                     if "INS AMT IN USD" in cell.text.replace("\xa0", " "):  # normalize NBSP
#                         parts = cell.text.split("-")
#                         if len(parts) > 1:
#                             value_text = parts[-1].strip()
#                             INS = remove_comma(value_text)
#                             try:
#                                 ins_amount = float(INS)
#                                 total_amount -= ins_amount
#                                 ins_yes = True
#                                 for p in cell.paragraphs:
#                                     for run in p.runs:
#                                         run.text = f"INS AMT IN USD - {ins_amount * 0.6:.2f}"  # reduce 40%
#                             except ValueError:
#                                 print(f"Skipping non-numeric INS amount: {INS}")

#                 if not fob_yes:
#                     if 'FOB Value IN USD -' in cell.text:
#                         x = cell.text.split('-')
#                         space = sum(1 for xx in x[-1] if xx == ' ') if len(x) > 1 else 0
#                         x = x[-1].lstrip()
#                         for paragraph in cell.paragraphs:
#                             for run in paragraph.runs:
#                                 run.text = run.text.replace(cell.text, f'FOB Value IN USD -{(space * " ")+formatAmount(total_amount)}')
#                         fob_yes = True

# def processPDF(input_pdf_path, output_pdf_path):
#     error, _ = Licence()
#     if error != "":
#         print("Error:", error)
#     else:
#         try:
#             doc_file_path = input_pdf_path.replace(".pdf", ".docx")
#             pdftodocx(input_pdf_path, doc_file_path)
#             font_change(doc_file_path)

#             doc = Document(doc_file_path)
#             read_tables_from_docx(doc)
            
#             doc.save(doc_file_path)
#             print(f"DOCX file updated and saved at: {doc_file_path}")

#             pythoncom.CoInitialize()
#             convert(doc_file_path, output_pdf_path)
#             os.remove(doc_file_path)
#             print(f"DONE: PDF saved at {output_pdf_path}")

#         except Exception as e:
#             print(f"ERROR: {e}")

# # Example usage
# input_pdf_path = r"D:\Python\Pdf_Manipulation\lateste_fob\2513200404 CI.pdf"  # Update as needed
# output_pdf_path = r"D:\Python\Pdf_Manipulation\lateste_fob\2513200404_updated.pdf"  # Update as needed

# processPDF(input_pdf_path, output_pdf_path)



import tkinter as tk
import webbrowser
from pdf2docx import Converter
from docx2pdf import convert
from num2words import num2words
from docx.shared import Pt
from pathlib import Path  # core python module
import os
import pythoncom
from tkinter import *
from tkinter import messagebox
from docx import Document
from tkinter import filedialog , messagebox
import webbrowser 
import sys
import time
import requests
import json
import locale
from invoice2data import extract_data
from invoice2data.extract.loader import read_templates


locale.setlocale(locale.LC_ALL, '')
def Licence():
        error=""
        message=""
        url = "http://aisagar.com/DeepBluePDFManipulation"
        payload = json.dumps({
        "token": "JUYHJ-LKJH12-IKJUYH-45OIUJ"
        })
        headers = {
        'Content-Type': 'application/json'
        }
        response = requests.request("POST", url, headers=headers, data=payload)    
        obj = json.loads(response.text)
        if "Error" in response.text and "Message" in response.text:
            error=obj["Error"]
            message=obj["Message"]
        else:
            error=response.text    
        return error,message


def reduce_invoice():
    global Licence
    error,_=Licence()
    if error!="":
        messagebox.showerror("showwarning", error)
    else:
        global root 
        root.destroy()

        def pdftodocx(pdf_file_path,doc_file_path):
            obj=Converter(pdf_file_path)
            obj.convert(doc_file_path)
            obj.close()

        def font_change(doc_file_path):
            def change_font_size(doc, target_text, new_font_size):
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if target_text in cell.text:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.size = Pt(new_font_size)

        # Replace 'your_invoice.docx' with the actual file path of your DOCX file
            #file_path = 'input/output1.docx'

            # Load the document
            doc = Document(doc_file_path)
            # Specify the target text and the new font size
            target_text = 'Port of Discharge:'
            new_font_size = 9
            # Call the function to change the font size
            change_font_size(doc, target_text, new_font_size)
            # Save the modified document
            modified_file_path = doc_file_path
            doc.save(modified_file_path)


        #change 
        def decrease_number_by_percent(number):
            decrease_percentage=40
            decrease_factor = 1 - (decrease_percentage / 100)
            decreased_number = number * decrease_factor
            return float(decreased_number)

        def remove_comma(num):
            return num.replace(",",'')

        def usd_change(amount):
        # amount = 29719.85
            # print(amount)
            a1=str(amount)
            q1,q2=a1.split('.')
            words1 = num2words(int(q1), lang='en').capitalize()
            words2 = num2words(int(q2), lang='en').capitalize()
            words1=words1.replace(" and",'')
            words1+=' dollars and '
            words2+=' cents only'
            word = words1.upper()+words2.upper()
            return word

        def formatAmount(amt):
            amt = locale.format_string("%.2f", amt, grouping=True)
            return str(amt)

        def hsn(doc):
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if 'HSN CODE:'  in cell.text :
                            string = str(cell.text)
                            string= string.replace('HSN CODE:','')
                            all = string.split()
                            return all

        def is_number(s):
            """Return True if s can be converted to float."""
            try:
                float(s.replace(",", ""))
                return True
            except ValueError:
                return False

        def decrease_number_by_percent(value):
            """Reduce the value by 40% (retain 60%)."""
            return value * 0.6

        def read_tables_from_docx(doc):
            total_qty = 0.0
            total_amount = 0.0
            total_steel = 0.0  # Accumulate decreased steel values
            frt_amount = 0.0
            ins_amount = 0.0
            frt_yes = False
            ins_yes = False
            fob_yes = False
            data_table_found = False
            total_row = None  # Initialize total_row outside the loop

            for table in doc.tables:
                for row_idx, row in enumerate(table.rows):
                    cells = row.cells
                    print("Cells:", [cell.text.strip() for cell in cells])  # Debug print

                    # Check if the table has enough columns for a data table
                    if len(cells) < 12:  # Minimum columns required (adjusted for multi-page case)
                        print(f"Skipping row {row_idx} due to insufficient columns: {len(cells)} columns")
                        continue

                    # Detect start of data table
                    if not data_table_found and cells[0].text.strip() == "S.N":
                        data_table_found = True
                        continue  # Skip header row

                    if data_table_found:
                        # Identify and store the total row
                        if "Total:" in cells[0].text:
                            total_row = row
                            continue

                        try:
                            # Qty (adjust index based on structure)
                            qty_text = cells[9].text.strip().replace(",", "") if len(cells) > 9 else "0"
                            qty = float(qty_text) if is_number(qty_text) else 0.0
                            total_qty += qty
                            print(f"Row {row_idx}: Qty = {qty}, Total Qty = {total_qty}")

                            # Rate in USD (col 16, adjust if needed)
                            rate_text = cells[16].text.strip().replace(",", "") if len(cells) > 16 else ""
                            if is_number(rate_text):
                                rate = float(rate_text)
                                decreased_rate = round(decrease_number_by_percent(rate), 6)
                                for p in cells[16].paragraphs:
                                    for run in p.runs:
                                        run.text = f"{decreased_rate:.6f}"
                                new_amount = round(qty * decreased_rate, 2)
                                total_amount += new_amount
                                for p in cells[17].paragraphs:
                                    for run in p.runs:
                                        run.text = formatAmount(new_amount)
                                print(f"Row {row_idx}: Rate = {rate} -> {decreased_rate}, New Amount = {new_amount}, Total Amount = {total_amount}")

                            # Rate of Steel Per KG (col 12)
                            rate_steel_text = cells[12].text.strip().replace(",", "") if len(cells) > 12 else ""
                            if is_number(rate_steel_text):
                                rate_steel = float(rate_steel_text)
                                decreased_rate_steel = round(decrease_number_by_percent(rate_steel), 6)
                                for p in cells[12].paragraphs:
                                    for run in p.runs:
                                        run.text = f"{decreased_rate_steel:.6f}"

                            # Total value of Steel content used (col 13)
                            steel_text = cells[13].text.strip().replace(",", "") if len(cells) > 13 else ""
                            if is_number(steel_text):
                                steel_val = float(steel_text)
                                decreased_steel = round(decrease_number_by_percent(steel_val), 2)
                                total_steel += decreased_steel  # Accumulate decreased steel value
                                for p in cells[13].paragraphs:
                                    for run in p.runs:
                                        run.text = formatAmount(decreased_steel)
                                print(f"Row {row_idx}: Steel Val = {steel_val} -> {decreased_steel}, Total Steel = {total_steel}")

                        except IndexError as e:
                            print(f"Error processing row {row_idx}: {e}")
                            continue

            # After processing all rows and tables, update the total amount row
            if total_row is not None:
                # Dynamically find the indices of numeric values in the total row
                numeric_values = [(i, float(cell.text.strip().replace(",", ""))) for i, cell in enumerate(total_row.cells) if is_number(cell.text.strip())]
                if len(numeric_values) >= 2:
                    amount_idx = numeric_values[-1][0]  # Last numeric value (Grand Total in USD)
                    steel_idx = numeric_values[-2][0]   # Second-to-last numeric value (Total value of Steel content used)
                else:
                    amount_idx = 17  # Fallback to default if not enough numeric values
                    steel_idx = 13

                # Debug the indices and current values
                print(f"Updating steel_idx = {steel_idx}, amount_idx = {amount_idx}")
                print(f"Current Steel Value: {total_row.cells[steel_idx].text.strip()}")
                print(f"Current Amount Value: {total_row.cells[amount_idx].text.strip()}")

                # Update Grand Total in USD
                has_paragraphs = False
                for p in total_row.cells[amount_idx].paragraphs:
                    has_paragraphs = True
                    for run in p.runs:
                        run.text = formatAmount(round(total_amount, 2))
                        print(f"Setting Grand Total in USD (col {amount_idx}) to {formatAmount(round(total_amount, 2))}")
                if not has_paragraphs:
                    print(f"No paragraphs found in Grand Total cell {amount_idx}")
                    # Fallback to direct text update if no paragraphs
                    total_row.cells[amount_idx].text = formatAmount(round(total_amount, 2))

                # Update Total value of Steel content used
                has_paragraphs = False
                formatted_total_steel = formatAmount(round(total_steel, 2))
                original_total_steel = total_row.cells[steel_idx].text.strip()
                for p in total_row.cells[steel_idx].paragraphs:
                    has_paragraphs = True
                    for run in p.runs:
                        run.text = run.text.replace(original_total_steel, formatted_total_steel)
                        print(f"Setting Total value of Steel content used (col {steel_idx}) to {formatted_total_steel}")
                if not has_paragraphs:
                    print(f"No paragraphs found in Steel Total cell {steel_idx}")
                    # Fallback to direct text update if no paragraphs
                    total_row.cells[steel_idx].text = formatted_total_steel

                # Print the cell text after updates
                print("Updated Total Row Cells:")
                print(f"Total value of Steel content used (col {steel_idx}): {total_row.cells[steel_idx].text.strip()}")
                print(f"Grand Total in USD (col {amount_idx}): {total_row.cells[amount_idx].text.strip()}")
        
            # Update other sections (unchanged)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if 'Amount in words: USD' in cell.text:
                            amount_words = cell.text.replace('Amount in words: USD ', '')
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.text = run.text.replace(amount_words, usd_change(total_amount))

                        if not frt_yes:    
                            if 'FRT AMT IN USD -' in cell.text:
                                x = cell.text.split('-')
                                space = sum(1 for xx in x[-1] if xx == ' ') if len(x) > 1 else 0
                                x = x[-1].lstrip()
                                FRT = remove_comma(x)
                                frt_amount = float(FRT)
                                frt_amount = decrease_number_by_percent(frt_amount)
                                total_amount -= frt_amount
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.text = run.text.replace(cell.text, f'FRT AMT IN USD -{(space * " ")+formatAmount(frt_amount)}')
                                frt_yes = True

                        if not ins_yes:
                            if "INS AMT IN USD" in cell.text.replace("\xa0", " "):  # normalize NBSP
                                parts = cell.text.split("-")
                                if len(parts) > 1:
                                    value_text = parts[-1].strip()
                                    INS = remove_comma(value_text)
                                    try:
                                        ins_amount = float(INS)
                                        total_amount -= ins_amount
                                        ins_yes = True
                                        for p in cell.paragraphs:
                                            for run in p.runs:
                                                run.text = f"INS AMT IN USD - {ins_amount * 0.6:.2f}"  # reduce 40%
                                    except ValueError:
                                        print(f"Skipping non-numeric INS amount: {INS}")

                        if not fob_yes:
                            if 'FOB Value IN USD -' in cell.text:
                                x = cell.text.split('-')
                                space = sum(1 for xx in x[-1] if xx == ' ') if len(x) > 1 else 0
                                x = x[-1].lstrip()
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.text = run.text.replace(cell.text, f'FOB Value IN USD -{(space * " ")+formatAmount(total_amount)}')
                                fob_yes = True

        def processPDF():
            error,_=Licence()
            if error!="":
                messagebox.showerror("showwarning", error) 
            else:
                input_Folder=filedialog.askdirectory()
                output_Folder="D:/Python/Pdf_Manipulation/lateste_fob/updated_fob/"
                count=0
                files = [file for file in os.listdir(input_Folder) if file.lower().endswith(".pdf")]
                for file in files:
                    try:
                        count=count+1
                        progress_label_var.set("Processing "+str(count)+"/"+str(len(files)))
                        root1.update()

                        writeLog("PROCESSING: "+ file)

                        doc_file_path="D:/Python/Pdf_Manipulation/lateste_fob/updated_fob/"+file.replace(".pdf","")+".docx"
                        pdf_path =os.path.join(input_Folder, file)
                        pdftodocx(pdf_path,doc_file_path)
                        font_change(doc_file_path)
                        
                        doc = Document(doc_file_path)
                        HSN= hsn(doc)
                        read_tables_from_docx(doc)
                        doc.save(doc_file_path)
                        ########

                        
                        pythoncom.CoInitialize()
                        sys.stderr=open("consoleoutput.log","w")
                        convert(doc_file_path,output_Folder)
                        os.remove(doc_file_path)
                        writeLog("DONE: "+ file)

                    except Exception as e:
                        writeLog("ERROR: "+ str(e))
                        
                # lb_done = Label(panel,text="DONE!") 
                # lb_done.grid(row=3, column=0, sticky="w", padx=5, pady=5)
                progress_label_var.set("DONE!")
                root1.update()    

        def writeLog(message):
            tm = time.strftime('%d-%b-%Y %H:%M:%S')
            txt_log.insert(END, tm + " - " + message + "\n")
            txt_log.yview(END)
            root1.update()

        def on_link_click(event):
            webbrowser.open("https://www.sagarinfotech.com")

        def Licence():
            error=""
            message=""
            url = "http://aisagar.com/DeepBluePDFManipulation"
            payload = json.dumps({
            "token": "JUYHJ-LKJH12-IKJUYH-45OIUJ"
            })
            headers = {
            'Content-Type': 'application/json'
            }
            response = requests.request("POST", url, headers=headers, data=payload)    
            obj = json.loads(response.text)
            if "Error" in response.text and "Message" in response.text:
                error=obj["Error"]
                message=obj["Message"]
            else:
                error=response.text    
            return error,message

        def showMessage():
            _,message=Licence()
            lbl_Message = Label(panel, text = message,bg=panel_color,fg = "red")
            lbl_Message.grid(row=0, column=0, sticky="w", padx=15, pady=5)

        # create root window
        root1 = Tk()
        root1.title("Welcome to AI World(SIPL)")
        root1.iconbitmap("content/SIPL.ico")
        # Set geometry (widthxheight)
        root1.geometry('400x460') 

        panel_color = '#E0E0E0'
        panel = Frame(root1, bg=panel_color, padx=5, pady=5)
        panel.grid(row=0, column=0, padx=5, pady=5, sticky='nsew')

        #adding a label to the root window
        Path("C:/Output_PDF").mkdir(parents=True, exist_ok=True)

        root1.after(5000, showMessage)

        lbl = Label(panel, text = "Output Folder: C:\Output_PDF",bg=panel_color)
        lbl.grid(row=1, column=0, sticky="w", padx=15, pady=5)

        # Create a button with specified font properties
        button_font = ("Arial", 10, "bold")
        btn = Button(panel, text = "Select Invoice Folder and Start Process",font=button_font ,bg="#fa8f16",fg = "#fff", command=processPDF)
        btn.grid(row=2, column=0, sticky="w", padx=15, pady=5)
        root1.resizable(width=False, height=False)

        progress_label_var = StringVar()
        progress_label_var.set(" ")
        label = Label(panel, textvariable=progress_label_var,bg=panel_color)
        label.grid(row=3, column=0, sticky="w", padx=5, pady=5)

        txt_log = Text(panel, height=19, width=58,font=("Helvetica", 8))
        txt_log.grid(row=4, column=0, sticky="w", padx=5, pady=5)

        scrollbar = Scrollbar(panel, command=txt_log.yview)
        scrollbar.grid(row=4, column=1, padx=0, pady=5,sticky="ns")
        txt_log.config(yscrollcommand=scrollbar.set)

        label_lnk = Label(panel, text="Developed By: Sagarinfotech.com", font=('Helvetica', 8, 'underline'),bg=panel_color, fg='blue', cursor='hand2')
        label_lnk.grid(row=5, column=0, sticky="w", padx=1, pady=1)
        label_lnk.bind("<Button-1>", on_link_click)

        # Execute Tkinter
        root1.mainloop()


# Create main window
if __name__=="__main__":
    ## if comm_invoice_fob.txt file exist then remove 
    
    #------
    
    def on_link_click(event):
        webbrowser.open("https://www.sagarinfotech.com")

    #error,_=Licence()
     
    def showMessage():
        _,message=Licence()
        lbl_Message = Label(root , text = message,bg=panel_color,fg = "red")
        lbl_Message.grid(row=0, column=0, sticky="w", padx=15, pady=5)

    # create root window
    root = Tk()
    root.title("Welcome to AI World(SIPL)")
    root.iconbitmap("content/SIPL.ico")
    # Set geometry (widthxheight) 
    root.geometry('450x185') 

    panel_color = '#E0E0E0'

    # panel = Frame(root, bg=panel_color, padx=5, pady=5)
    # panel.grid(row=0, column=0, padx=5, pady=5, sticky='nsew')

    #adding a label to the root window
  
    root.after(5000, showMessage)
    textff= tk.Label( root  , text=" ")
    textff.grid ( row = 1 , column= 0)

    # Create a button with specified font properties
    button_font = ("Arial", 10, "bold")
    btn = Button( root , text = "REDUCE INVOICE ",bg="lightblue", fg="black", padx=10, pady=5, relief="raised", font=("Helvetica", 12, "bold"), command=reduce_invoice)
    btn.grid(row=2, column=0, sticky="w", padx=15, pady=5)
    # btn1 = Button(root  , text = "VALIDATE AMOUNT ",bg="lightgreen", fg="black", padx=10, pady=5, relief="raised", font=("Helvetica", 12, "bold"), command=validate_amount)
    # btn1.grid(row=2, column=1, sticky="w", padx=15, pady=5)
    root.resizable(width=False, height=False)
    text1 = tk.Label(root , text = " ")
    text1.grid(row=3,column= 0 )
    text2 = tk.Label(root , text = " ")
    text2.grid(row=4,column= 0 )
    text3 = tk.Label(root , text = " ")
    text3.grid(row=5,column= 0 )

    label_lnk = Label(root, text="Developed By: Sagarinfotech.com", font=('Helvetica', 8, 'underline'), bg=panel_color, fg='blue', cursor='hand2')
    label_lnk.grid(row=6, column=0, padx=1, pady=1, columnspan=2, sticky='sw')
    label_lnk.bind("<Button-1>", on_link_click)

    # Execute Tkinter
    root.mainloop()
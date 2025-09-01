import requests
import json
from pdf2docx import Converter
from docx2pdf import convert
from num2words import num2words
from docx.shared import Pt
from pathlib import Path
import os
import pythoncom
import time
import locale
from docx import Document

locale.setlocale(locale.LC_ALL, '')

def Licence():
    error = ""
    message = ""
    url = "http://aisagar.com/DeepBluePDFManipulation"
    payload = json.dumps({
        "token": "JUYHJ-LKJH12-IKJUYH-45OIUJ"
    })
    headers = {
        'Content-Type': 'application/json'
    }
    response = requests.request("POST", url, headers=headers, data=payload)    
    obj = json.loads(response.text)
    if "Error" in response.text and "Message" in response.text:
        error = obj["Error"]
        message = obj["Message"]
    else:
        error = response.text    
    return error, message

def pdftodocx(pdf_file_path, doc_file_path):
    obj = Converter(pdf_file_path)
    obj.convert(doc_file_path)
    obj.close()

def font_change(doc_file_path):
    def change_font_size(doc, target_text, new_font_size):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if target_text in cell.text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(new_font_size)

    doc = Document(doc_file_path)
    target_text = 'Port of Discharge:'
    new_font_size = 9
    change_font_size(doc, target_text, new_font_size)
    doc.save(doc_file_path)

def decrease_number_by_percent(number):
    decrease_percentage = 60
    decrease_factor = 1 - (decrease_percentage / 100)
    decreased_number = number * decrease_factor
    return float(decreased_number)

def remove_comma(num):
    return num.replace(",", '')

def usd_change(amount):
    a1 = str(amount)
    q1, q2 = a1.split('.')
    words1 = num2words(int(q1), lang='en').capitalize()
    words2 = num2words(int(q2), lang='en').capitalize()
    words1 = words1.replace(" and", '')
    words1 += ' dollars and '
    words2 += ' cents only'
    word = words1.upper() + words2.upper()
    return word

# def formatAmount(amt):
#     # Truncate to 2 decimal places to avoid rounding
#     amt = int(amt * 100) / 100.0
#     amt = locale.format_string("%.2f", amt, grouping=True)
#     return str(amt)


def formatAmount(amt):
    # Explicitly round to 2 decimal places
    amt = round(amt, 2)
    # Use Python's string formatting to avoid locale issues
    formatted_amt = f"{amt:.2f}"  # Outputs "769.31"
    print(f"formatAmount input: {amt}, output: {formatted_amt}")  # Debug print
    return formatted_amt

def hsn(doc):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if 'HSN CODE:' in cell.text:
                    string = str(cell.text)
                    string = string.replace('HSN CODE:', '')
                    all = string.split()
                    return all

def read_tables_from_docx(docx, HSN):
    ii = 0 
    doc = docx
    process = 0 
    starting = False
    total_next = 0 
    total_amount = float()
    total_yes = False
    words_yes = False
    frt_yes = False
    ins_yes = False
    fob_yes = False
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text in HSN:
                    process = 1 
                if process:
                    if cell.text not in HSN:
                        if ii == 1:
                            quantity = cell.text 
                        elif ii == 2:
                            decreased_number_rateusd = decrease_number_by_percent(float(remove_comma(cell.text)))
                            decreased_number_rateusd = round(decreased_number_rateusd, 6)
                            value_derived_amtusd = float(decreased_number_rateusd * int(remove_comma(quantity)))
                            print("value_derived_amtusd:", value_derived_amtusd)
                            value_derived_amtusd = round(value_derived_amtusd, 6)  # From second code
                            print("value_derived_amtusd1:", value_derived_amtusd)
                            # total_amount += value_derived_amtusd
                            total_amount += round(value_derived_amtusd, 2)
                            print("total_amount:",total_amount)

                            str_decreased_number_rateusd = str(decreased_number_rateusd)
                            orginal_rate = cell.text
                            print("orginal_rate:", orginal_rate) 
                            str_decreased_number_rateusd = str(decreased_number_rateusd)
                            if len(orginal_rate.split(".")) > 1:
                                if len(orginal_rate.split(".")[1]) > len(str_decreased_number_rateusd.split(".")[1]):
                                    str_decreased_number_rateusd = str(decreased_number_rateusd) + ("0" * (len(orginal_rate.split(".")[1]) - len(str_decreased_number_rateusd.split(".")[1]) - 1))
                            
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.text = run.text.replace(cell.text, str_decreased_number_rateusd + " ")

                        # elif ii == 3:
                        #     for paragraph in cell.paragraphs:
                        #         for run in paragraph.runs:
                        #             run.text = run.text.replace(cell.text, "  " + formatAmount(value_derived_amtusd))
                        #             print("Updated Run:", run.text)
                                    
                        #     ii = 0 
                        #     process = 0 
                        #     continue
                        elif ii == 3:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    print(f"Original cell text in ii == 3: {cell.text}")  # Debug print
                                    formatted_value = formatAmount(value_derived_amtusd)
                                    print("formatted_value:",formatted_value)
                                    run.text = run.text.replace(cell.text, "  " + formatted_value)
                                    print("Updated Run:", run.text)
                                    
                            ii = 0 
                            process = 0 
                            continue

                        ii += 1





                if 'Total:' in cell.text:
                    starting = True
                if starting:
                    if len(cell.text) == 0:
                        total_next = 1 
                    if not total_yes:
                        if total_next:
                            if len(cell.text) != 0:
                                total_amount = round(total_amount, 2)
                                print("total_amount1:",total_amount)
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.text = run.text.replace(cell.text, formatAmount(total_amount) + " ")
                                        total_yes = True
                                starting = False
                if 'Amount in words: USD' in cell.text:
                    amount_words = cell.text
                    amount_words = amount_words.replace('Amount in words: USD ', '')
                    total_amount = round(total_amount, 2)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.text = run.text.replace(amount_words, usd_change(total_amount))

                if not frt_yes:    
                    if 'FRT AMT IN USD -' in cell.text:
                        x = cell.text.split('-')
                        space = 0
                        for xx in x[-1]:
                            if xx == ' ':
                                space += 1 
                            else:
                                break
                        x = x[-1].lstrip()
                        FRT = x 
                        FRT = remove_comma(FRT)
                        FRT = float(FRT)
                        FRT = decrease_number_by_percent(FRT)
                        total_amount -= FRT
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.text = run.text.replace(cell.text, f'FRT AMT IN USD -{(space * " ")+formatAmount(FRT)}')
                                frt_yes = True

                if not ins_yes:
                    if 'INS AMT IN USD -' in cell.text:
                        x = cell.text.split('-')
                        x = x[-1].lstrip()
                        INS = x 
                        INS = remove_comma(INS)
                        INS = float(INS)
                        total_amount -= INS
                        ins_yes = True
                
                if not fob_yes:
                    if 'FOB Value IN USD -' in cell.text:
                        x = cell.text.split('-')
                        space = 0
                        for xx in x[-1]:
                            if xx == ' ':
                                space += 1 
                            else:
                                break
                        x = x[-1].lstrip()
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.text = run.text.replace(cell.text, f'FOB Value IN USD -{(space * " ")+formatAmount(total_amount)}')
                                fob_yes = True

def processPDF(input_pdf_path, output_pdf_path):
    error, _ = Licence()
    if error != "":
        print("Error:", error)
    else:
        try:
            # Convert PDF to DOCX
            doc_file_path = input_pdf_path.replace(".pdf", ".docx")
            pdftodocx(input_pdf_path, doc_file_path)
            font_change(doc_file_path)

            # Load DOCX file
            doc = Document(doc_file_path)
            HSN = hsn(doc)
            read_tables_from_docx(doc, HSN)
            
            # Save DOCX and verify changes
            doc.save(doc_file_path)
            print(f"DOCX file updated and saved at: {doc_file_path}")

            # Convert DOCX to PDF
            pythoncom.CoInitialize()
            convert(doc_file_path, output_pdf_path)
            os.remove(doc_file_path)
            print(f"DONE: PDF saved at {output_pdf_path}")

        except Exception as e:
            print(f"ERROR: {e}")

# Example usage
input_pdf_path = r"D:\Python\Pdf_Manipulation\input_folder\COMM 1162.pdf"  # Update with the path to your input PDF
output_pdf_path = r"D:\Python\Pdf_Manipulation\output_folder\COMM_1162_updated.pdf"  # Update with the desired path for the output PDF

processPDF(input_pdf_path, output_pdf_path)
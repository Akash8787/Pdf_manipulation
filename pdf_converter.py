import tkinter as tk
import webbrowser
from pdf2docx import Converter
from docx2pdf import convert
from num2words import num2words
from docx.shared import Pt
from pathlib import Path
import os
import pythoncom
from tkinter import filedialog, messagebox
from docx import Document
import sys
import time
import locale
import pdfplumber
import re
from pdfminer.high_level import extract_text
import math
import logging
import requests
import json

# Setup logging for validate_amount
logging.basicConfig(filename='invoice_processing.log', level=logging.ERROR)

locale.setlocale(locale.LC_ALL, '')

def Licence():
    error = ""
    message = ""
    url = "http://aisagar.com/DeepBluePDFManipulation"
    payload = json.dumps({
        "token": "JUYHJ-LKJH12-IKJUYH-45OIUJ"
    })
    headers = {
        'Content-Type': 'application/json'
    }
    try:
        response = requests.request("POST", url, headers=headers, data=payload)
        obj = json.loads(response.text)
        if "Error" in obj and "Message" in obj:
            error = obj["Error"]
            message = obj["Message"]
        else:
            error = response.text
    except Exception as e:
        error = f"License check failed: {str(e)}"
    return error, message

def reduce_invoice():
    global root
    error, _ = Licence()
    if error != "":
        messagebox.showerror("License Error", error)
        return
    root.destroy()

    def pdftodocx(pdf_file_path, doc_file_path):
        obj = Converter(pdf_file_path)
        obj.convert(doc_file_path)
        obj.close()

    def font_change(doc_file_path):
        def change_font_size(doc, target_text, new_font_size):
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if target_text in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(new_font_size)

        doc = Document(doc_file_path)
        target_text = 'Port of Discharge:'
        new_font_size = 9
        change_font_size(doc, target_text, new_font_size)
        modified_file_path = doc_file_path
        doc.save(modified_file_path)

    def decrease_number_by_percent(number):
        decrease_percentage = 60
        decrease_factor = 1 - (decrease_percentage / 100)
        decreased_number = number * decrease_factor
        return float(decreased_number)

    def remove_comma(num):
        return num.replace(",", '')

    def usd_change(amount):
        a1 = str(amount)
        q1, q2 = a1.split('.')
        words1 = num2words(int(q1), lang='en').capitalize()
        words2 = num2words(int(q2), lang='en').capitalize()
        words1 = words1.replace(" and", '')
        words1 += ' dollars and '
        words2 += ' cents only'
        word = words1.upper() + words2.upper()
        return word

    def formatAmount(amt):
        amt = locale.format_string("%.2f", amt, grouping=True)
        return str(amt)

    def hsn(doc):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if 'HSN CODE:' in cell.text:
                        string = str(cell.text)
                        string = string.replace('HSN CODE:', '')
                        all = string.split()
                        return all

    def read_tables_from_docx(doc, HSN):
        ii = 0
        process = 0
        starting = False
        total_next = 0
        total_amount = float()
        total_yes = False
        frt_yes = False
        ins_yes = False
        fob_yes = False

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text in HSN:
                        process = 1
                    if process:
                        if cell.text not in HSN:
                            if ii == 1:
                                quantity = cell.text
                            elif ii == 2:
                                decreased_number_rateusd = decrease_number_by_percent(float(remove_comma(cell.text)))
                                decreased_number_rateusd = round(decreased_number_rateusd, 6)
                                value_derived_amtusd = float(decreased_number_rateusd * int(remove_comma(quantity)))
                                value_derived_amtusd = round(value_derived_amtusd, 6)
                                total_amount += value_derived_amtusd

                                str_decreased_number_rateusd = str(decreased_number_rateusd)
                                orginal_rate = cell.text
                                if len(orginal_rate.split(".")) > 1:
                                    if len(orginal_rate.split(".")[1]) > len(str_decreased_number_rateusd.split(".")[1]):
                                        str_decreased_number_rateusd = str(decreased_number_rateusd) + ("0" * (len(orginal_rate.split(".")[1]) - len(str_decreased_number_rateusd.split(".")[1]) - 1))

                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.text = run.text.replace(cell.text, str_decreased_number_rateusd + " ")

                            elif ii == 3:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.text = run.text.replace(cell.text, "  " + formatAmount(value_derived_amtusd))
                                ii = 0
                                process = 0
                                continue

                            ii += 1
                    if 'Total:' in cell.text:
                        starting = True
                    if starting:
                        if len(cell.text) == 0:
                            total_next = 1
                        if not total_yes:
                            if total_next:
                                if len(cell.text) != 0:
                                    total_amount = round(total_amount, 2)
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.text = run.text.replace(cell.text, formatAmount(total_amount) + " ")
                                            total_yes = True
                                    starting = False
                    if 'Amount in words: USD' in cell.text:
                        amount_words = cell.text
                        amount_words = amount_words.replace('Amount in words: USD ', '')
                        total_amount = round(total_amount, 2)
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.text = run.text.replace(amount_words, usd_change(total_amount))

                    if not frt_yes:
                        if 'FRT AMT IN USD -' in cell.text:
                            x = cell.text.split('-')
                            space = 0
                            for xx in x[-1]:
                                if xx == ' ':
                                    space += 1
                                else:
                                    break
                            x = x[-1].lstrip()
                            FRT = x
                            FRT = remove_comma(FRT)
                            FRT = float(FRT)
                            FRT = decrease_number_by_percent(FRT)
                            total_amount -= FRT
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.text = run.text.replace(cell.text, f'FRT AMT IN USD -{(space * " ") + formatAmount(FRT)}')
                                    frt_yes = True

                    if not ins_yes:
                        if 'INS AMT IN USD -' in cell.text:
                            x = cell.text.split('-')
                            x = x[-1].lstrip()
                            INS = x
                            INS = remove_comma(INS)
                            INS = float(INS)
                            total_amount -= INS
                            ins_yes = True

                    if not fob_yes:
                        if 'FOB Value IN USD -' in cell.text:
                            x = cell.text.split('-')
                            space = 0
                            for xx in x[-1]:
                                if xx == ' ':
                                    space += 1
                                else:
                                    break
                            x = x[-1].lstrip()
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.text = run.text.replace(cell.text, f'FOB Value IN USD -{(space * " ") + formatAmount(total_amount)}')
                                    fob_yes = True

    def processPDF():
        error, _ = Licence()
        if error != "":
            messagebox.showerror("License Error", error)
            return
        input_Folder = filedialog.askdirectory()
        output_Folder = "C:\\Output_PDF\\"
        count = 0
        files = [file for file in os.listdir(input_Folder) if file.lower().endswith(".pdf")]
        for file in files:
            try:
                count += 1
                progress_label_var.set("Processing " + str(count) + "/" + str(len(files)))
                root1.update()

                writeLog("PROCESSING: " + file)

                doc_file_path = "C:\\Output_PDF\\" + file.replace(".pdf", "") + ".docx"
                pdf_path = os.path.join(input_Folder, file)
                pdftodocx(pdf_path, doc_file_path)
                font_change(doc_file_path)

                doc = Document(doc_file_path)
                HSN = hsn(doc)
                read_tables_from_docx(doc, HSN)
                doc.save(doc_file_path)

                pythoncom.CoInitialize()
                sys.stderr = open("consoleoutput.log", "w")
                convert(doc_file_path, output_Folder)
                os.remove(doc_file_path)
                writeLog("DONE: " + file)

            except Exception as e:
                writeLog("ERROR: " + str(e))

        progress_label_var.set("DONE!")
        root1.update()

    def writeLog(message):
        tm = time.strftime('%d-%b-%Y %H:%M:%S')
        txt_log.insert(tk.END, tm + " - " + message + "\n")
        txt_log.yview(tk.END)
        root1.update()

    def on_link_click(event):
        webbrowser.open("https://www.sagarinfotech.com")

    def showMessage():
        _, message = Licence()
        lbl_Message = tk.Label(panel, text=message, bg=panel_color, fg="red")
        lbl_Message.grid(row=0, column=0, sticky="w", padx=15, pady=5)

    # Create root window - reduce_invoice GUI
    root1 = tk.Tk()
    root1.title("Welcome to AI World(SIPL)")
    root1.iconbitmap("content/SIPL.ico")
    root1.geometry('400x460')

    panel_color = '#E0E0E0'
    panel = tk.Frame(root1, bg=panel_color, padx=5, pady=5)
    panel.grid(row=0, column=0, padx=5, pady=5, sticky='nsew')

    Path("C:/Output_PDF").mkdir(parents=True, exist_ok=True)

    root1.after(5000, showMessage)

    lbl = tk.Label(panel, text="Output Folder: C:\\Output_PDF", bg=panel_color)
    lbl.grid(row=1, column=0, sticky="w", padx=15, pady=5)

    button_font = ("Arial", 10, "bold")
    btn = tk.Button(panel, text="Select Invoice Folder and Start Process", font=button_font, bg="#fa8f16", fg="#fff", command=processPDF)
    btn.grid(row=2, column=0, sticky="w", padx=15, pady=5)
    root1.resizable(width=False, height=False)

    progress_label_var = tk.StringVar()
    progress_label_var.set(" ")
    label = tk.Label(panel, textvariable=progress_label_var, bg=panel_color)
    label.grid(row=3, column=0, sticky="w", padx=5, pady=5)

    txt_log = tk.Text(panel, height=19, width=58, font=("Helvetica", 8))
    txt_log.grid(row=4, column=0, sticky="w", padx=5, pady=5)

    scrollbar = tk.Scrollbar(panel, command=txt_log.yview)
    scrollbar.grid(row=4, column=1, padx=0, pady=5, sticky="ns")
    txt_log.config(yscrollcommand=scrollbar.set)

    label_lnk = tk.Label(panel, text="Developed By: Sagarinfotech.com", font=('Helvetica', 8, 'underline'), bg=panel_color, fg='blue', cursor='hand2')
    label_lnk.grid(row=5, column=0, sticky="w", padx=1, pady=1)
    label_lnk.bind("<Button-1>", on_link_click)

    root1.mainloop()

def validate_amount():
    global root
    error, _ = Licence()
    if error != "":
        messagebox.showerror("License Error", error)
        return
    root.destroy()

    def extract_text_lines(pdf_file):
        all_lines = []
        try:
            with pdfplumber.open(pdf_file) as pdf:
                for i, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text()
                    if text:
                        page_lines = text.strip().split('\n')
                        print(f"\n--- Page {i} ---\n")
                        for line in page_lines:
                            print(line)
                        all_lines.extend(page_lines)
                    else:
                        print(f"\n--- Page {i} ---\n[Empty Page]")
            if not all_lines:
                raise ValueError("Extracted text is empty. The PDF might be scanned or corrupted.")
            return all_lines
        except Exception as e:
            logging.error(f"Error reading PDF: {str(e)}")
            raise

    def extract_invoice_fields(text_lines, source_file):
        global total_mpf, total_hmf, total_duty
        extracted = []
        current_item = {}
        # total_mpf = ""
        # total_hmf = ""
        # total_duty = ""

        for i, line in enumerate(text_lines):
            line = line.strip()

            # Extract Line number (e.g., "001") from any line starting with 3 digits (but NOT 499 or 501)
            if re.match(r'^\d{3}\b', line) and not line.startswith(("499", "501")):
                if current_item:
                    extracted.append(current_item)

                current_item = {
                    "Line": line.strip().split()[0],
                    "Extracted Value": "",
                    "Duty1": "",
                    "Duty2": "",
                    "MPF": "",
                    "HMF": "",
                    "Invoice Number": "",
                    "Invoice Value": "",
                    "Entered Value": "",
                    "Source File": source_file
                }

            # Extract Duty1 from line with tariff code, weight, and FREE (e.g., "9903.01.28 2,094 KG FREE $0.00")
            if re.search(r'^\d{4}\.\d{2}\.\d{2}\s+[\d,]+\s+KG\b', line):

                match = re.search(r"\$\d[\d,]*\.\d{2}", line)
                if match:
                    current_item["Duty1"] = match.group()
                    # print(current_item["Duty1"])
            # Normalize line
            clean_line = re.sub(r"[^\x00-\x7F]+", " ", line)  # Remove non-ASCII characters

            # Unified condition
            if (
                (re.search(r'\d{4}\.\d{2}\.\d{4}', line) and '%' in line)
                or ("KG" in line and "$" in line)
            ):
                # dollar_matches = re.findall(r"\$\d{1,3}(?:,\d{3})*(?:\.\d{2})?\b", clean_line)
                dollar_matches = re.findall(r"\$\d[\d,]*(?:\.\d{2})?\b", clean_line)

                
                # print("Dollar_Values:", dollar_matches)

                if len(dollar_matches) >= 2:
                    current_item["Extracted Value"] = dollar_matches[0]
                    current_item["Duty2"] = dollar_matches[1]
                elif len(dollar_matches) == 1:
                    # Fallback logic depends on which case it is
                    if "KG" in line:
                        current_item["Extracted Value"] = dollar_matches[0]
                    else:
                        current_item["Duty2"] = dollar_matches[0]



            # Extract MPF
            if line.startswith("499 -") and "%" in line:
                match = re.search(r"\$\d[\d,]*\.\d{2}", line)
                if match:
                    current_item["MPF"] = match.group()

            # Extract HMF
            if line.startswith("501 -") and "%" in line:
                match = re.search(r"\$\d[\d,]*\.\d{2}", line)
                if match:
                    current_item["HMF"] = match.group()


            if line.startswith("499 - MPF"):
                match = re.search(r"\$\d[\d,]*\.\d{2}", line)
                if match:
                    total_mpf = match.group()
                    print("Total MPF:", total_mpf)

            if line.startswith("501 - HMF"):
                match = re.search(r"\$\d[\d,]*\.\d{2}", line)
                if match:
                    total_hmf = match.group()
                    print("Total HMF:", total_hmf)
            # Look for total duty amount that comes AFTER the line containing "Ascertained Duty"
            if "Ascertained Duty" in line:
                # Look ahead up to 3 lines after current line to find a standalone dollar value
                for j in range(1, 4):
                    if i + j < len(text_lines):
                        next_line = text_lines[i + j].strip()
                        match = re.match(r"^\$\d[\d,]*\.\d{2}$", next_line)
                        if match:
                            total_duty = match.group()
                            print("Total Duty:", total_duty)
                            break



            # Totals for Invoice line
            if "Totals for Invoice" in line:
                next_line = text_lines[i + 1].strip() if i + 1 < len(text_lines) else ""
                invoice_info = re.findall(r"\d{10}", next_line)
                values = re.findall(r"\d[\d,]*\.\d{2}\sUSD", next_line)
                if invoice_info:
                    current_item["Invoice Number"] = invoice_info[0]
                if len(values) >= 2:
                    current_item["Invoice Value"] = values[0]
                    current_item["Entered Value"] = values[1]

        # Append the last item
        if current_item:
            extracted.append(current_item)
        print(extracted)

        return extracted

    def clean(var, add_currency=False, add_usd=False):
        if not isinstance(var, str):
            return float(var) if var else 0.0
        var = var.replace('$', '').replace(',', '').replace('%', '').replace('USD', '').strip()
        if var.lower() in ('no', 'n', 'free', ''):
            return 0.0
        try:
            cleaned_value = float(var)
            if add_currency:
                return f"${cleaned_value:.2f}"
            if add_usd:
                return f"{cleaned_value:.2f} USD"
            return cleaned_value
        except ValueError:
            return 0.0

    def clean_exporter_details(text):
        cleaned = re.sub(r'^\w+:\n', '', text.strip())
        cleaned = re.sub(r'^\s*\w+[: ]*\n', '', cleaned)
        cleaned = re.sub(r'\n\nInvoice No:.*', '', cleaned, flags=re.DOTALL)
        return cleaned.strip()

    def extract_invoice_folder_data(invoice_folder):
        all_json_data = []
        if not os.path.exists(invoice_folder):
            return all_json_data, f"Invoice folder not found: {invoice_folder}"

        for filename in os.listdir(invoice_folder):
            if filename.lower().endswith(".pdf"):
                pdf_path = os.path.join(invoice_folder, filename)
                try:
                    text = extract_text(pdf_path)
                except Exception as e:
                    print(f"Error extracting text from {filename}: {e}")
                    continue

                start_marker_exporter = "Manufacture & Exporter :\n"
                end_marker_exporter = "\nIEC Code:"
                start_index_exporter = text.find(start_marker_exporter) + len(start_marker_exporter)
                end_index_exporter = text.find(end_marker_exporter)
                if start_index_exporter == -1 or end_index_exporter == -1:
                    exporter_details = "Manufacture & Exporter details not found."
                else:
                    exporter_details = clean_exporter_details(text[start_index_exporter:end_index_exporter])

                start_marker_invoice = "Invoice No: "
                end_marker_invoice = "\n"
                start_index_invoice = text.find(start_marker_invoice) + len(start_marker_invoice)
                end_index_invoice = text.find(end_marker_invoice, start_index_invoice)
                if start_index_invoice == -1 or end_index_invoice == -1:
                    invoice_no = "Invoice No. not found."
                else:
                    invoice_no = text[start_index_invoice:end_index_invoice].strip()

                start_marker_fob = "FOB Value IN USD - "
                end_marker_fob = "\n"
                start_index_fob = text.find(start_marker_fob) + len(start_marker_fob)
                end_index_fob = text.find(end_marker_fob, start_index_fob)
                if start_index_fob == -1 or end_index_fob == -1:
                    fob_value = "FOB Value not found."
                else:
                    fob_value = text[start_index_fob:end_index_fob].strip()

                data = {
                    "pdf_file": filename,
                    "Manufacture & Exporter": exporter_details,
                    "Invoice No.": invoice_no,
                    "FOB Value IN USD": fob_value
                }

                if "not found" in exporter_details:
                    continue

                try:
                    fob_numeric = float(fob_value.replace(',', '').strip())
                    data["FOB Value IN USD"] = fob_numeric
                except ValueError:
                    print(f"Skipped {filename} due to invalid FOB value: {fob_value}")
                    continue

                all_json_data.append(data)
                print(f"Processed {filename}")

        return all_json_data, ""


    # def process_pdf(pdf_file):
    #     error, _ = Licence()
    #     if error != "":
    #         return f"❌ License Error: {error}"
    #     try:
    #         output = f"Processing main PDF: {pdf_file}\n"
    #         invoice_folder = entry_invoice_folder.get().strip() or r"D:\Python\Amount_Validation\59therror\FOB INV. DETAILS US ATRIM STAR CONT. NO. 15"

    #         lines = extract_text_lines(pdf_file)
    #         results = extract_invoice_fields(lines, os.path.basename(pdf_file))

    #         total_added_duty = 0.0
    #         total_added_mpf = 0.0
    #         total_added_hmf = 0.0

    #         for item in results:
    #             total_added_duty += clean(item.get("Duty1", 0))
    #             total_added_duty += clean(item.get("Duty2", 0))
    #             total_added_mpf += clean(item.get("MPF", 0))
    #             total_added_hmf += clean(item.get("HMF", 0))
    #         print("total_added_duty:",total_added_duty)
    #         print("total_added_mpf:",total_added_mpf)
    #         print("total_added_hmf:",total_added_hmf)

    #         total_duty_val = float(clean(total_duty))
    #         total_mpf_val = float(clean(total_mpf))
    #         total_hmf_val = float(clean(total_hmf))
    #         print("total_duty_val:",total_duty)
    #         print("total_mpf_val:",total_mpf)
    #         print("total_hmf_val:",total_hmf)
    #         tolerance = 1e-6

    #         output += "\n--- Comparison Results for Totals ---\n"
    #         output += "✅ Total Duty matches.\n" if math.isclose(total_added_duty, total_duty_val, abs_tol=tolerance) \
    #             else f"❌ Total Duty mismatch! Expected: {total_duty_val:.2f}, Found: {total_added_duty:.2f}\n"
    #         output += "✅ Total MPF matches.\n" if math.isclose(total_added_mpf, total_mpf_val, abs_tol=tolerance) \
    #             else f"❌ Total MPF mismatch! Expected: {total_mpf_val:.2f}, Found: {total_added_mpf:.2f}\n"
    #         output += "✅ Total HMF matches.\n" if math.isclose(total_added_hmf, total_hmf_val, abs_tol=tolerance) \
    #             else f"❌ Total HMF mismatch! Expected: {total_hmf_val:.2f}, Found: {total_added_hmf:.2f}\n"




    #         all_json_data, error_msg = extract_invoice_folder_data(invoice_folder)
    #         if error_msg:
    #             output += f"❌ {error_msg}\n"
    #             return output

    #         output += "\n--- Invoice Comparison Results ---\n"
    #         invoice_fob_list = [(item["Invoice No."], item["FOB Value IN USD"]) for item in all_json_data]
    #         matched_invoices = set()

    #         for invoice, fob in invoice_fob_list:
    #             found = False
    #             for item in results:
    #                 if str(item["Invoice Number"]).strip() == str(invoice).strip():
    #                     found = True
    #                     matched_invoices.add(invoice)
    #                     entered_value = clean(item.get("Entered Value", "0"))
    #                     invoice_value = clean(item.get("Invoice Value", "0"))
    #                     if (math.isclose(fob, entered_value, rel_tol=1e-3) and
    #                         math.isclose(fob, invoice_value, rel_tol=1e-3)):
    #                         output += f"for Invoice No. {invoice} OK ✅ \n"
    #                     else:
    #                         output += f"❌ Value mismatch for Invoice No. {invoice}: FOB={fob}, Entered={entered_value}, Invoice={invoice_value}\n"
    #                     break

    #         all_invoices_ok = True

    #         invoice_folder_invoices = {inv for inv, _ in invoice_fob_list}
    #         for item in results:
    #             invoice_no = str(item.get("Invoice Number", "")).strip()
    #             if invoice_no and invoice_no not in invoice_folder_invoices:
    #                 output += f"📂 Invoice {invoice_no} from ENTRY SUMMARY pdf has no matching file in Invoice Folder.\n"
    #                 all_invoices_ok = False

    #         entry_summary_invoices = {str(item.get("Invoice Number", "")).strip() for item in results}
    #         for inv, _ in invoice_fob_list:
    #             if inv not in entry_summary_invoices:
    #                 output += f"⚠️ Invoice No. {inv} not found in ENTRY SUMMARY.\n"
    #                 all_invoices_ok = False

    #         if invoice_fob_list or results:
    #             if all_invoices_ok:
    #                 output += "\nOverall OK ✅\n"
    #             else:
    #                 output += "\n Overall NOT OK ❌\n"

    #         return output


    #         # invoice_folder_invoices = {inv for inv, _ in invoice_fob_list}
    #         # entry_summary_invoices = {str(item.get("Invoice Number", "")).strip() for item in results}

    #         # all_invoices_ok = True
    #         # output = "\n--- Invoice Comparison Results ---\n"

    #         # # For each invoice in ENTRY SUMMARY, check if it exists in the folder
    #         # for item in results:
    #         #     invoice_no = str(item.get("Invoice Number", "")).strip()
    #         #     if invoice_no:
    #         #         if invoice_no in invoice_folder_invoices:
    #         #             output += f"for Invoice No. {invoice_no} OK ✅\n"
    #         #         else:
    #         #             output += f"📂 Invoice {invoice_no} from ENTRY SUMMARY PDF has no matching file in Invoice Folder.\n"
    #         #             all_invoices_ok = False

    #         # # Optional: Only warn for invoices in the folder that are expected in summary
    #         # # In other words, don't flag extra files
    #         # # So this check is now removed completely or restricted

    #         # # Final summary
    #         # if invoice_fob_list or results:
    #         #     output += "\nOverall OK ✅\n" if all_invoices_ok else "\nOverall NOT OK ❌\n"

    #         # return output
        # except Exception as e:
        #     logging.error(f"Error processing {pdf_file}: {str(e)}", exc_info=True)
        #     return f"❌ Error processing {pdf_file}: {str(e)}"



    def process_pdf(pdf_file):
        try:
            # License check
            error, _ = Licence()
            if error:
                return f"❌ License Error: {error}"

            output = f"Processing main PDF: {pdf_file}\n"
            invoice_folder = entry_invoice_folder.get().strip() or r"D:\Python\Amount_Validation\59therror\FOB INV. DETAILS US ATRIM STAR CONT. NO. 15"

            # Extract lines and fields
            lines = extract_text_lines(pdf_file)
            results = extract_invoice_fields(lines, os.path.basename(pdf_file))

            total_added_duty = 0.0
            total_added_mpf = 0.0
            total_added_hmf = 0.0

            for item in results:
                total_added_duty += clean(item.get("Duty1", 0))
                total_added_duty += clean(item.get("Duty2", 0))
                total_added_mpf += clean(item.get("MPF", 0))
                total_added_hmf += clean(item.get("HMF", 0))

            # Debug print
            print("Extracted Totals -> Duty:", total_added_duty, "MPF:", total_added_mpf, "HMF:", total_added_hmf)

            # Ensure total_* values are defined
            try:
                total_duty_val = float(clean(total_duty))
                total_mpf_val = float(clean(total_mpf))
                total_hmf_val = float(clean(total_hmf))
            except NameError as ne:
                return f"❌ One or more total values (total_duty, total_mpf, total_hmf) are not defined: {ne}"
            except Exception as e:
                return f"❌ Error processing total values: {e}"

            print("Expected Totals -> Duty:", total_duty_val, "MPF:", total_mpf_val, "HMF:", total_hmf_val)

            tolerance = 1e-6
            output += "\n--- Comparison Results for Totals ---\n"
            output += "✅ Total Duty matches.\n" if math.isclose(total_added_duty, total_duty_val, abs_tol=tolerance) \
                else f"❌ Total Duty mismatch! Expected: {total_duty_val:.2f}, Found: {total_added_duty:.2f}\n"
            output += "✅ Total MPF matches.\n" if math.isclose(total_added_mpf, total_mpf_val, abs_tol=tolerance) \
                else f"❌ Total MPF mismatch! Expected: {total_mpf_val:.2f}, Found: {total_added_mpf:.2f}\n"
            output += "✅ Total HMF matches.\n" if math.isclose(total_added_hmf, total_hmf_val, abs_tol=tolerance) \
                else f"❌ Total HMF mismatch! Expected: {total_hmf_val:.2f}, Found: {total_added_hmf:.2f}\n"

            # Extract invoice data from folder
            all_json_data, error_msg = extract_invoice_folder_data(invoice_folder)
            if error_msg:
                output += f"❌ {error_msg}\n"
                return output

            output += "\n--- Invoice Comparison Results ---\n"
            invoice_fob_list = [(item["Invoice No."], item["FOB Value IN USD"]) for item in all_json_data]
            matched_invoices = set()

            for invoice, fob in invoice_fob_list:
                found = False
                for item in results:
                    if str(item.get("Invoice Number", "")).strip() == str(invoice).strip():
                        found = True
                        matched_invoices.add(invoice)
                        entered_value = clean(item.get("Entered Value", "0"))
                        invoice_value = clean(item.get("Invoice Value", "0"))

                        if (math.isclose(fob, entered_value, rel_tol=1e-3) and
                            math.isclose(fob, invoice_value, rel_tol=1e-3)):
                            output += f"for Invoice No. {invoice} OK ✅\n"
                        else:
                            output += f"❌ Value mismatch for Invoice No. {invoice}: FOB={fob}, Entered={entered_value}, Invoice={invoice_value}\n"
                        break

            invoice_folder_invoices = {inv for inv, _ in invoice_fob_list}
            entry_summary_invoices = {str(item.get("Invoice Number", "")).strip() for item in results}

            all_invoices_ok = True

            # Check if all Entry Summary invoices are found in the Invoice Folder
            for invoice_no in entry_summary_invoices:
                if invoice_no and invoice_no not in invoice_folder_invoices:
                    output += f"📂 Invoice {invoice_no} from ENTRY SUMMARY pdf has no matching file in Invoice Folder.\n"
                    all_invoices_ok = False

            # ✅ Don't warn about extra invoices in folder if all summary invoices are found
            # i.e., only show ⚠️ if an invoice from the folder is missing in the summary AND summary wasn't fully matched
            if not entry_summary_invoices.issubset(invoice_folder_invoices):
                for inv, _ in invoice_fob_list:
                    if inv not in entry_summary_invoices:
                        output += f"⚠️ Invoice No. {inv} not found in ENTRY SUMMARY.\n"
                        all_invoices_ok = False

            if invoice_fob_list or results:
                if all_invoices_ok:
                    output += "\nOverall OK ✅\n"
                else:
                    output += "\n Overall NOT OK ❌\n"


        except Exception as e:
            return f"❌ Unexpected error in process_pdf: {e}"





    def browse_pdf():
        file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
        if file_path:
            entry_pdf_path.delete(0, tk.END)
            entry_pdf_path.insert(0, file_path)

    def browse_invoice_folder():
        folder_path = filedialog.askdirectory()
        if folder_path:
            entry_invoice_folder.delete(0, tk.END)
            entry_invoice_folder.insert(0, folder_path)

    def run_processing():
        error, _ = Licence()
        if error != "":
            messagebox.showerror("License Error", error)
            return
        pdf_file = entry_pdf_path.get().strip()
        if not pdf_file:
            messagebox.showerror("Error", "Please select a main PDF file.")
            return
        result = process_pdf(pdf_file)
        output_text.config(state=tk.NORMAL)
        output_text.delete(1.0, tk.END)
        output_text.insert(tk.END, result)
        output_text.config(state=tk.DISABLED)

    # Create root window for validate_amount
    root = tk.Tk()
    root.title("PDF Invoice Processor")
    root.geometry("800x600")

    frame_top = tk.Frame(root)
    frame_top.pack(pady=10)

    tk.Label(frame_top, text="Select Entry PDF File:").pack(side=tk.LEFT, padx=5)
    entry_pdf_path = tk.Entry(frame_top, width=60)
    entry_pdf_path.pack(side=tk.LEFT, padx=5)
    btn_browse_pdf = tk.Button(frame_top, text="Browse", command=browse_pdf)
    btn_browse_pdf.pack(side=tk.LEFT, padx=5)

    frame_folder = tk.Frame(root)
    frame_folder.pack(pady=10)
    tk.Label(frame_folder, text="Select Invoice Folder:").pack(side=tk.LEFT, padx=5)
    entry_invoice_folder = tk.Entry(frame_folder, width=60)
    entry_invoice_folder.pack(side=tk.LEFT, padx=5)
    btn_browse_folder = tk.Button(frame_folder, text="Browse", command=browse_invoice_folder)
    btn_browse_folder.pack(side=tk.LEFT, padx=5)

    btn_run = tk.Button(root, text="Run Processing", command=run_processing, bg="green", fg="white", width=20)
    btn_run.pack(pady=10)

    output_text = tk.Text(root, wrap=tk.WORD, height=25, width=100, state=tk.DISABLED)
    output_text.pack(padx=10, pady=10)

    def on_link_click(event):
        webbrowser.open("https://www.sagarinfotech.com")

    label_lnk = tk.Label(root, text="Developed By: Sagarinfotech.com", font=('Helvetica', 8, 'underline'), fg='blue', cursor='hand2')
    label_lnk.pack(side="left", fill="x")
    label_lnk.bind("<Button-1>", on_link_click)

    root.mainloop()

# Main GUI
if __name__ == "__main__":
    def on_link_click(event):
        webbrowser.open("https://www.sagarinfotech.com")

    def showMessage():
        _, message = Licence()
        lbl_Message = tk.Label(root, text=message, bg=panel_color, fg="red")
        lbl_Message.grid(row=0, column=0, sticky="w", padx=15, pady=5, columnspan=2)

    root = tk.Tk()
    root.title("Welcome to AI World(SIPL)")
    root.iconbitmap("content/SIPL.ico")
    root.geometry('450x185')

    panel_color = '#E0E0E0'

    root.after(5000, showMessage)

    textff = tk.Label(root, text=" ")
    textff.grid(row=1, column=0)

    button_font = ("Arial", 10, "bold")
    btn = tk.Button(root, text="REDUCE INVOICE", bg="lightblue", fg="black", padx=10, pady=5, relief="raised", font=("Helvetica", 12, "bold"), command=reduce_invoice)
    btn.grid(row=2, column=0, sticky="w", padx=15, pady=5)
    btn1 = tk.Button(root, text="VALIDATE AMOUNT", bg="lightgreen", fg="black", padx=10, pady=5, relief="raised", font=("Helvetica", 12, "bold"), command=validate_amount)
    btn1.grid(row=2, column=1, sticky="w", padx=15, pady=5)
    root.resizable(width=False, height=False)

    text1 = tk.Label(root, text=" ")
    text1.grid(row=3, column=0)
    text2 = tk.Label(root, text=" ")
    text2.grid(row=4, column=0)
    text3 = tk.Label(root, text=" ")
    text3.grid(row=5, column=0)

    label_lnk = tk.Label(root, text="Developed By: Sagarinfotech.com", font=('Helvetica', 8, 'underline'), bg=panel_color, fg='blue', cursor='hand2')
    label_lnk.grid(row=6, column=0, padx=1, pady=1, columnspan=2, sticky='sw')
    label_lnk.bind("<Button-1>", on_link_click)

    root.mainloop()
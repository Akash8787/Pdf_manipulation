import tkinter as tk
import webbrowser
from pdf2docx import Converter
from docx2pdf import convert
from num2words import num2words
from docx.shared import Pt
from pathlib import Path  # core python module
import os
import pythoncom
from tkinter import *
from tkinter import messagebox
from docx import Document
from tkinter import filedialog , messagebox
import webbrowser 
import sys
import time
import requests
import json
import locale
from invoice2data import extract_data
from invoice2data.extract.loader import read_templates
import re


locale.setlocale(locale.LC_ALL, '')
def Licence():
        error=""
        message=""
        url = "http://aisagar.com/DeepBluePDFManipulation"
        payload = json.dumps({
        "token": "JUYHJ-LKJH12-IKJUYH-45OIUJ"
        })
        headers = {
        'Content-Type': 'application/json'
        }
        response = requests.request("POST", url, headers=headers, data=payload)    
        obj = json.loads(response.text)
        if "Error" in response.text and "Message" in response.text:
            error=obj["Error"]
            message=obj["Message"]
        else:
            error=response.text    
        return error,message


def reduce_invoice():
    global Licence
    error,_= Licence()
    if error!="":
        messagebox.showerror("showwarning", error)
    else:
        global root 
        root.destroy()

        def pdftodocx(pdf_file_path,doc_file_path):
            obj=Converter(pdf_file_path)
            obj.convert(doc_file_path)
            obj.close()

        def font_change(doc_file_path):
            def change_font_size(doc, target_text, new_font_size):
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if target_text in cell.text:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.size = Pt(new_font_size)

        # Replace 'your_invoice.docx' with the actual file path of your DOCX file
            #file_path = 'input/output1.docx'

            # Load the document
            doc = Document(doc_file_path)
            # Specify the target text and the new font size
            target_text = 'Port of Discharge:'
            new_font_size = 9
            # Call the function to change the font size
            change_font_size(doc, target_text, new_font_size)
            # Save the modified document
            modified_file_path = doc_file_path
            doc.save(modified_file_path)


        #change 
        def decrease_number_by_percent(number):
            decrease_percentage=60
            decrease_factor = 1 - (decrease_percentage / 100)
            decreased_number = number * decrease_factor
            return float(decreased_number)

        def remove_comma(num):
            return num.replace(",",'')

        def usd_change(amount):
        # amount = 29719.85
            # print(amount)
            a1=str(amount)
            q1,q2=a1.split('.')
            words1 = num2words(int(q1), lang='en').capitalize()
            words2 = num2words(int(q2), lang='en').capitalize()
            words1=words1.replace(" and",'')
            words1+=' dollars and '
            words2+=' cents only'
            word = words1.upper()+words2.upper()
            return word

        def formatAmount(amt):
            amt = locale.format_string("%.2f", amt, grouping=True)
            return str(amt)

        def hsn(doc):
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if 'HSN CODE:'  in cell.text :
                            string = str(cell.text)
                            string= string.replace('HSN CODE:','')
                            all = string.split()
                            return all

        def read_tables_from_docx(docx , HSN):
            ii= 0 
            doc = docx
            process =  0 
            starting = False
            total_next =  0 
            total_amount = float()
            total_yes  = False
            words_yes = False
            frt_yes = False
            ins_yes = False
            fob_yes = False
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text in HSN:
                            process= 1 
                        if process:
                            if cell.text not in HSN:
                                if ii==1:
                                    # print('quantity = ', cell.text)
                                    quantity = cell.text 
                                elif ii==2:
                                    # print('rate = ', cell.text) 
                                    
                                    decreased_number_rateusd=decrease_number_by_percent(float(remove_comma(cell.text)))
                                    

                                    decreased_number_rateusd=round(decreased_number_rateusd,6)
                                    value_derived_amtusd = float(decreased_number_rateusd*int(remove_comma(quantity)))
                                    value_derived_amtusd=round(value_derived_amtusd,6)
                                    total_amount+=value_derived_amtusd

                                    str_decreased_number_rateusd=str(decreased_number_rateusd)
                                    orginal_rate = cell.text 
                                    str_decreased_number_rateusd=str(decreased_number_rateusd)
                                    if len(orginal_rate.split("."))>1:
                                        if len(orginal_rate.split(".")[1])>len(str_decreased_number_rateusd.split(".")[1]):
                                            str_decreased_number_rateusd=str(decreased_number_rateusd)+("0"*(len(orginal_rate.split(".")[1])-len(str_decreased_number_rateusd.split(".")[1])-1))
                                        
                                    for paragraph in cell.paragraphs:
                                        for run in  paragraph.runs:
                                            run.text = run.text.replace(cell.text ,str_decreased_number_rateusd+" ")

                                elif ii==3:
                                    # print('amount  = ', cell.text)
                                    for paragraph in cell.paragraphs:
                                        for run in  paragraph.runs:

                                            run.text = run.text.replace(cell.text,"  "+formatAmount(value_derived_amtusd)  )
                                    ii= 0 
                                    process = 0 
                                    continue

                                ii+=1
                                # input( ' = ')
                        if 'Total:' in cell.text:
                            starting = True
                        if starting:
                            if len(cell.text )==0:
                                total_next = 1 
                            if not total_yes:
                                if total_next:
                                    if len(cell.text)!= 0:
                                        total_amount=round(total_amount,2)

                                        # print( 'total --',  cell.text)
                                        for paragraph in cell.paragraphs:
                                                for run in  paragraph.runs:
                                                    run.text = run.text.replace(cell.text,formatAmount(total_amount)+" " )
                                                    total_yes = True

                                        starting = False
                        if 'Amount in words: USD' in cell.text :
                            # print('word = ', cell.text)
                            amount_words=cell.text
                            amount_words=amount_words.replace('Amount in words: USD ','')
                            total_amount=round(total_amount,2)

                            for paragraph in cell.paragraphs:
                                        for run in  paragraph.runs:
                                            run.text = run.text.replace(amount_words ,usd_change(total_amount) )

                        if not frt_yes:    
                            if 'FRT AMT IN USD -' in cell.text :
                                
                                x = cell.text.split('-')
                                space = 0
                                for xx in x[-1]:
                                    if xx ==' ':
                                        space+=1 
                                    else:
                                        break

                                x = x[-1].lstrip()

                                # print(  x     ) 
                                FRT = x 
                                FRT = remove_comma(FRT)
                                FRT=float(FRT)
                                FRT=decrease_number_by_percent(FRT)
                                total_amount-=FRT
                                for paragraph in cell.paragraphs:
                                    for run in  paragraph.runs:
                                        run.text = run.text.replace(cell.text ,f'FRT AMT IN USD -{(space*" ")+formatAmount(FRT)}')
                                        frt_yes = True

                        if not ins_yes:
                            if 'INS AMT IN USD -' in cell.text :

                                x = cell.text.split('-')
                                x = x[-1].lstrip()
                                # print(  x)
                                INS = x 
                                INS=remove_comma(INS)
                                INS = float(INS)
                                total_amount-=INS
                                ins_yes = True
                        
                        if not fob_yes:
                            if 'FOB Value IN USD -' in cell.text :
                                x = cell.text.split('-')
                                space = 0
                                for xx in x[-1]:
                                    if xx ==' ':
                                        space+=1 
                                    else:
                                        break

                                x = x[-1].lstrip()
                                for paragraph in cell.paragraphs:
                                    for run in  paragraph.runs:
                                        run.text = run.text.replace(cell.text, f'FOB Value IN USD -{(space*" ")+formatAmount(total_amount)}')
                                        fob_yes = True

                                # print(  x)


        def processPDF():
            error,_=Licence()
            if error!="":
                messagebox.showerror("showwarning", error) 
            else:
                input_Folder=filedialog.askdirectory()
                output_Folder="D:\Python\Pdf_Manipulation\output_folder\\"
                count=0
                files = [file for file in os.listdir(input_Folder) if file.lower().endswith(".pdf")]
                for file in files:
                    try:
                        count=count+1
                        progress_label_var.set("Processing "+str(count)+"/"+str(len(files)))
                        root1.update()

                        writeLog("PROCESSING: "+ file)

                        doc_file_path="D:\Python\Pdf_Manipulation\output_folder\\"+file.replace(".pdf","")+".docx"
                        pdf_path =os.path.join(input_Folder, file)
                        pdftodocx(pdf_path,doc_file_path)
                        font_change(doc_file_path)
                        
                        doc = Document(doc_file_path)
                        HSN= hsn(doc)
                        read_tables_from_docx(doc , HSN)
                        doc.save(doc_file_path)
                        ########

                        
                        pythoncom.CoInitialize()
                        sys.stderr=open("consoleoutput.log","w")
                        convert(doc_file_path,output_Folder)
                        os.remove(doc_file_path)
                        writeLog("DONE: "+ file)

                    except Exception as e:
                        writeLog("ERROR: "+ str(e))
                        
                # lb_done = Label(panel,text="DONE!") 
                # lb_done.grid(row=3, column=0, sticky="w", padx=5, pady=5)
                progress_label_var.set("DONE!")
                root1.update()    

        def writeLog(message):
            tm = time.strftime('%d-%b-%Y %H:%M:%S')
            txt_log.insert(END, tm + " - " + message + "\n")
            txt_log.yview(END)
            root1.update()

        def on_link_click(event):
            webbrowser.open("https://www.sagarinfotech.com")

        def Licence():
            error=""
            message=""
            url = "http://aisagar.com/DeepBluePDFManipulation"
            payload = json.dumps({
            "token": "JUYHJ-LKJH12-IKJUYH-45OIUJ"
            })
            headers = {
            'Content-Type': 'application/json'
            }
            response = requests.request("POST", url, headers=headers, data=payload)    
            obj = json.loads(response.text)
            if "Error" in response.text and "Message" in response.text:
                error=obj["Error"]
                message=obj["Message"]
            else:
                error=response.text    
            return error,message

        def showMessage():
            _,message=Licence()
            lbl_Message = Label(panel, text = message,bg=panel_color,fg = "red")
            lbl_Message.grid(row=0, column=0, sticky="w", padx=15, pady=5)

        # create root window
        root1 = Tk()
        root1.title("Welcome to AI World(SIPL)")
        root1.iconbitmap("content/SIPL.ico")
        # Set geometry (widthxheight)
        root1.geometry('400x460') 

        panel_color = '#E0E0E0'
        panel = Frame(root1, bg=panel_color, padx=5, pady=5)
        panel.grid(row=0, column=0, padx=5, pady=5, sticky='nsew')

        #adding a label to the root window
        Path("C:/Output_PDF").mkdir(parents=True, exist_ok=True)

        root1.after(5000, showMessage)

        lbl = Label(panel, text = "Output Folder: C:\Output_PDF",bg=panel_color)
        lbl.grid(row=1, column=0, sticky="w", padx=15, pady=5)

        # Create a button with specified font properties
        button_font = ("Arial", 10, "bold")
        btn = Button(panel, text = "Select Invoice Folder and Start Process",font=button_font ,bg="#fa8f16",fg = "#fff", command=processPDF)
        btn.grid(row=2, column=0, sticky="w", padx=15, pady=5)
        root1.resizable(width=False, height=False)

        progress_label_var = StringVar()
        progress_label_var.set(" ")
        label = Label(panel, textvariable=progress_label_var,bg=panel_color)
        label.grid(row=3, column=0, sticky="w", padx=5, pady=5)

        txt_log = Text(panel, height=19, width=58,font=("Helvetica", 8))
        txt_log.grid(row=4, column=0, sticky="w", padx=5, pady=5)

        scrollbar = Scrollbar(panel, command=txt_log.yview)
        scrollbar.grid(row=4, column=1, padx=0, pady=5,sticky="ns")
        txt_log.config(yscrollcommand=scrollbar.set)

        label_lnk = Label(panel, text="Developed By: Sagarinfotech.com", font=('Helvetica', 8, 'underline'),bg=panel_color, fg='blue', cursor='hand2')
        label_lnk.grid(row=5, column=0, sticky="w", padx=1, pady=1)
        label_lnk.bind("<Button-1>", on_link_click)

        # Execute Tkinter
        root1.mainloop()


   
def validate_amount():

    def writeLog(message):
            tm = time.strftime('%d-%b-%Y %H:%M:%S')
            text_box.insert(END, tm + " - " + message + "\n")
            text_box.yview(END)
            root.update()

    global Licence
    error,_=Licence()

    if error!="":
        messagebox.showerror("showwarning", error)
    
    else: ### If no any error raise then run all functionalities
        global root 
        root.destroy()

        ## function for taking  EKH pdf paths  console files   
        console_file_selected = False
        pdf_folder_selected = False

        def browse_file(entry):
            global browse_file_var, console_file_selected
            browse_file_var=""
            filename = filedialog.askopenfilename()
            entry.delete(0, tk.END)
            entry.insert(0, filename)
            browse_file_var = filename
            console_file_selected = True              # Set flag to true when console file is selected

        def list_image_paths(folder_path):
            image_extensions = ['.pdf', '.PDF']  # Add more extensions if needed
            image_paths = []

            for rootfor, dirs, files in os.walk(folder_path):
                for file in files:
                    _, ext = os.path.splitext(file)
                    if ext.lower() in image_extensions:
                        image_paths.append(os.path.join(rootfor, file))

            return image_paths

        def clean_comm_txtfile():
            '''
            Removing duplicates from COMM txt file
            '''
            try:
                unique_lines = set()

                with open('comm_invoice_fob.txt', 'r') as file:
                    for line in file:
                        unique_lines.add(line.strip())

                # Write the unique lines back to the file
                with open('comm_invoice_fob.txt', 'w') as file:
                    for line in unique_lines:
                        file.write(line + '\n')
            except Exception as e:
                writeLog('Error while removing duplicates from Invoice FOB txt ' + str(e))

        def Adding_invoice_and_FOB(entry):
            '''
            Adding invoice value and FOB value from folders
            '''
            if browse_file_var!="":
                console_file_selected=True 

            global pdf_folder_selected,foldername
            try:
                f = open('comm_invoice_fob.txt', 'w')
                f.close()

                foldername = filedialog.askdirectory()
                entry.delete(0, tk.END)
                entry.insert(0, foldername)
                templates = read_templates("Template2")
                
                folder_path = foldername
                foldername_var = foldername
                pythoncom.CoInitialize()
                sys.stderr = open("consoleoutput.log", "w")
                for i in list_image_paths(folder_path):
                    result = extract_data(i, templates=templates)

                    if isinstance(result, dict):  # Check if result is a dictionary
                        with open("comm_invoice_fob.txt", "a") as file:
                            try:
                                invoice_number = str(result['invoice_number']).strip()
                                invoice_number_copy = ''
                                for n in invoice_number:
                                    if n.isdigit():
                                        invoice_number_copy += n

                                invoice_number = invoice_number_copy
                                fob_value = str(result['FOB Value IN USD -']).split()[-1].strip()

                                file.write(invoice_number + " " + fob_value + "\n")

                                text_box.insert(tk.END, "Added invoice number " + invoice_number + ", FOB " + fob_value + "\n")
                            except Exception as e:
                                writeLog("ERROR IN Adding_invoice_and_FOB: " + str(e))
                    else:
                        print("Data extraction failed for:", i)

                clean_comm_txtfile()
                text_box.insert(tk.END, "            " + "\n")
                pdf_folder_selected = True  # Set flag to true when PDF folder is selected
            except Exception as e:
                writeLog("ERROR IN Adding_invoice_and_FOB: " + str(e))

        def on_link_click(event):
            webbrowser.open("https://www.sagarinfotech.com")



        def StartProcess   ():
            global console_file_selected , pdf_folder_selected
            '''
            validating all data 
            '''

            if not console_file_selected:
                messagebox.showerror("Error", "Please select the console file first.")
                return
            if not pdf_folder_selected:
                messagebox.showerror("Error", "Please select the PDF folder.")
                return
           
            #removing reason.txt if already exist 
            if os.path.exists("reason.txt"):
                os.remove("reason.txt")
            
            #stored variable where all data stored 
            stored={} 
            pythoncom.CoInitialize()
            sys.stderr=open("consoleoutput.log","w")
            templates= read_templates("Template/")
            # print("Template:",templates)
            try:
                #result1 variable where data extract are   MPF,HMF,Total Entered Value , Total Other Fees,Duty,Other,Total from pdf through YML(regex)
                result1 = extract_data(browse_file_var,templates=templates) 
                print("Result1:",result1)
            except Exception as e :
                writeLog('Error while fetching  data except invoices  '+str(e))

            ## result variable where cleaned data stored 
            result={}
           
            def cleaning_data(result1):
                try:
                    nonlocal result
                    if 'A.' in result1['aa']:
                    
                        result1['aa']=  result1['aa'][0:  str(result1['aa']).index("A.")]
                        result['MPF']= float(result1['aa'].split("$")[-1].replace(",",'')) 
                        print("result['MPF']1:",result['MPF'])
                    else:
                    
                        result['MPF']= float(result1['aa'].split("$")[-1].replace(",",'')) 
                        print("result['MPF']2:",result['MPF'])
                    if 'A.' in result1['bb']:
                        result1['bb']=  result1['bb'][0:  str(result1['bb']).index("A.")]
                        result1['bb']=  result1['bb'][ str(result1['bb']).index("$") +1: ]
                        result['HMF'] = float(result1['bb'].split("$")[0].replace(",",''))
                        result['Total Entered Value'] = float(result1['bb'].split("$")[-1].replace(",",'') )
                    else:
                        result1['bb']=  result1['bb'][ str(result1['bb']).index("$") +1: ]
                        result['HMF'] = float(result1['bb'].split("$")[0].replace(",",''))
                        result['Total Entered Value'] = float(result1['bb'].split("$")[-1].replace(",",'') )
                    result['Duty'] = float(result1['cc'].split('$')[-1].replace(",",''))
                    # print("result['Duty']:",result['Duty'])
                    result['Total Other Fees']= float(result1['dd'].split("\n")[-1].replace("$",'').replace(",",''))
                    if "I" in result1['ee']:
                        result1['ee']=(result1['ee'].split("\n"))
                        result['Other'] = float(result1['ee'][0].split("$")[-1].replace("\r",'').replace("\n",'').replace(",",''))
                    else:
                        result1['ee']=(result1['ee'].split("\n"))
                        result['Other'] = float(result1['ee'][-1].split("$")[-1].replace("\r",'').replace("\n",'').replace(",",''))
                    result['Total']= float(result1['ff'].split("$")[-1].replace("\r",'').replace("\n",'').replace(",",''))
                except Exception as e :
                    writeLog('Error while cleaning data except invoices '+str(e))
           
            cleaning_data(result1)
            # def clean_check_block(raw_text):
            #         """Clean a single check block text to start from HS Code line."""
            #         # Split into lines
            #         lines = raw_text.splitlines()

            #         # Get the line number from the very first line
            #         line_number = lines[0].split()[0] if lines else ""

            #         # Find where the HS Code (8708.xx.xxxx) line starts
            #         try:
            #             start_idx = next(i for i, line in enumerate(lines) if re.search(r'\b8708\.\d{2}\.\d{4}\b', line))
            #             # Rebuild cleaned block: starting from HS Code, but keep line number attached
            #             cleaned_lines = [f" {line_number} " + lines[start_idx]] + lines[start_idx + 1:]
            #             cleaned_text = "\r\n".join(cleaned_lines)
            #             return cleaned_text
            #         except StopIteration:
            #             # If no 8708... line is found, return the original
            #             return raw_text     
            # temp1 = read_templates('Template3/')
            # # print("Template3:",temp1)
            # # try:
            # #     ## result2 variable where   Invoices details data extracted
            # #     result2 = extract_data(browse_file_var,templates=temp1)
            # #     print("Result2:",result2)
            # #     # print("--------------------------------")
            # # except Exception as e:
            # #     writeLog('Error while fetching all invoices details '+str(e))


            # try:
            #     # Extract invoice data
            #     result2 = extract_data(browse_file_var, templates=temp1)
            #     # print("Original Result2:", result2)

            #     # Check if 'check' field exists and clean it
            #     if result2 and 'check2' in result2:
            #         cleaned_checks = []
            #         for block in result2['check2']:
            #             cleaned_block = clean_check_block(block)
            #             cleaned_checks.append(cleaned_block)
            #         result2['check2'] = cleaned_checks  # Update with cleaned blocks

            #     # print("\nCleaned Result2:", result2)

            # except Exception as e:
            #     print('Error while fetching all invoices details: ' + str(e))
            def clean_check_block(raw_text):
                """Clean a single check block text to start from HS Code line and return removed portion."""
                lines = raw_text.splitlines()
                line_number = lines[0].split()[0] if lines else ""

                try:
                    # Find HS code line
                    start_idx = next(i for i, line in enumerate(lines) if re.search(r'\b8708\.\d{2}\.\d{4}\b', line))

                    # Cleaned block starts from HS code
                    cleaned_lines = [f" {line_number} " + lines[start_idx]] + lines[start_idx + 1:]
                    cleaned_text = "\r\n".join(cleaned_lines)

                    # Removed block is everything before HS code
                    removed_lines = lines[:start_idx]
                    removed_text = "\r\n".join(removed_lines)

                    return cleaned_text, removed_text
                except StopIteration:
                    return raw_text, ""  # Return original block, no removed section

            # Load templates
            temp1 = read_templates('Template3/')
            # browse_file_var = r"D:\Python\Pdf_Manipulation\input_folder\55therror.PDF"

            try:
                result2 = extract_data(browse_file_var, templates=temp1)
                # print("Original Result2:", result2)

                removed_texts = []  # To hold removed content

                if result2 and 'check2' in result2:
                    cleaned_checks = []
                    for block in result2['check2']:
                        cleaned_block, removed_block = clean_check_block(block)
                        cleaned_checks.append(cleaned_block)
                        removed_texts.append(removed_block)  # Store removed part

                    result2['check2'] = cleaned_checks
                    result2['removed_texts'] = removed_texts  # Add removed list to result2 for reference

                # print("\nCleaned Result2:", result2)


                # dollar_lines = []

                # Extract only the dollar amounts like $356.90, remove the '$', and convert to float
                dollar_values = [
                    float(match.replace(",", "").replace("$", ""))
                    for text in removed_texts
                    for match in re.findall(r'\$\d{1,3}(?:,\d{3})?\.\d{2}', text)
                ]

                print("Extracted Dollar Values:", dollar_values)
            except Exception as e:
                print('Error while fetching all invoices details: ' + str(e))

            

            stored_latest = {}
            
            def clean(var):
                    var = var.replace('$','').replace(',','').replace('%','').replace('USD','')
                    try:
                        if var!= 'NO' or var!= 'N':
                            var = float( var)
                    except ValueError:
                        pass
                    return var


            def handle_sep(result2):
                nonlocal stored_latest, clean
                
                try:
                    if type(result2['separatedboxes'])==str:
                        all_stored = []
                        text = str(result2['separatedboxes']).split("\n")
                        key = text[0].split()[0]
                        key = int(key)

                        second = text[1][text[1].find(  'N' )  : ].split()
                        for i in range(len(second)):
                            second[i]=clean(second[i])
                        
                        all_stored.extend( second[1:] )
                        forth = text[4].replace('499 - Merchandise Processing Fee','').split()
                        for i in range(len(forth)):
                            forth[i] = clean(forth[i])
                        all_stored.extend(forth)
                        fifth  = text[5].replace('501 - Harbor Maintenance Fee','').split()
                        for i in range(len(fifth)):
                            fifth[i] = clean(fifth[i])
                        all_stored.extend(fifth)
                        all_stored=[x for x in all_stored if x!='']
                        print("fifth1:",all_stored)

                        sixth = text[-1].replace('USD','').split()
                        for i in range(len(sixth)):
                            sixth[i] = clean(  sixth[i] ) 
                        all_stored.extend( sixth)


                        all_stored[7] = int(all_stored[7])
                        stored_latest[str(key)] = all_stored
                        # print("stored_latest:",stored_latest)
                        
                        
                    elif type(result2['separatedboxes'])==list:
                        for var in result2['separatedboxes']:
                            all_stored = []
                            text = str(var).split("\n")
                            key = text[0].split()[0]
                            key = int(key)
                            second = text[1][text[1].find(  'N' )  : ].split()
                            for i in range(len(second)):
                                second[i]=clean(second[i])
                            
                            all_stored.extend( second[1:] )
                            forth = text[4].replace('499 - Merchandise Processing Fee','').split()
                            for i in range(len(forth)):
                                forth[i] = clean(forth[i])
                            all_stored.extend(forth)
                            fifth  = text[5].replace('501 - Harbor Maintenance Fee','').split()
                            for i in range(len(fifth)):
                                fifth[i] = clean(fifth[i])
                            all_stored.extend(fifth)
                            all_stored=[x for x in all_stored if x!='']
                            print("fifth2:",all_stored)

                            sixth = text[-1].replace('USD','').split()
                            for i in range(len(sixth)):
                                sixth[i] = clean(  sixth[i] ) 
                            all_stored.extend( sixth)
                            all_stored[7] = int(all_stored[7])
                            stored_latest[str(key)]  = all_stored
                            # print("stored_latest2:",stored_latest)
                except KeyError :
                    try:
                        stored_update = {}
                        if type(result2['separatedboxes2'])==str:
                            all_stored = []
                            text = str(result2['separatedboxes2']).split("\n")
                            key = text[0].split()[0]
                            key = int(key)

                            second = text[0][text[0].find(  'N' )  : ].split()
                            for i in range(len(second)):
                                second[i]=clean(second[i])
                            
                            all_stored.extend( second[1:] )
                            forth = text[3].replace('499 - Merchandise Processing Fee','').split()
                            for i in range(len(forth)):
                                forth[i] = clean(forth[i])
                            all_stored.extend(forth)
                            fifth  = text[4].replace('501 - Harbor Maintenance Fee','').split()
                            for i in range(len(fifth)):
                                fifth[i] = clean(fifth[i])
                            all_stored.extend(fifth)
                            all_stored=[x for x in all_stored if x!='']
                            print("fifth3:",all_stored)

                            sixth = text[-1].replace('USD','').split()
                            for i in range(len(sixth)):
                                sixth[i] = clean(  sixth[i] ) 
                            all_stored.extend( sixth)


                            all_stored[7] = int(all_stored[7])
                            stored_update[str(key)] = all_stored
                            
                        
                        elif type(result2['separatedboxes2'])==list:
                            for var in result2['separatedboxes2']:
                                all_stored = []
                                text = str(var).split("\n")
                                key = text[0].split()[0]
                                key = int(key)
                                second = text[0][text[0].find(  'N' )  : ].split()
                                for i in range(len(second)):
                                    second[i]=clean(second[i])
                                
                                all_stored.extend( second[1:] )
                                forth = text[3].replace('499 - Merchandise Processing Fee','').split()
                                for i in range(len(forth)):
                                    forth[i] = clean(forth[i])
                                all_stored.extend(forth)
                                fifth  = text[4].replace('501 - Harbor Maintenance Fee','').split()
                                for i in range(len(fifth)):
                                    fifth[i] = clean(fifth[i])
                                all_stored.extend(fifth)
                                all_stored=[x for x in all_stored if x!='']
                                print("fifth4:",all_stored)

                                sixth = text[-1].replace('USD','').split()
                                for i in range(len(sixth)):
                                    sixth[i] = clean(  sixth[i] ) 
                                all_stored.extend( sixth)
                                all_stored[7] = int(all_stored[7])
                                stored_update[str(key)]  = all_stored
                        stored.update(stored_update)
                    except :

                        pass            


                    
                

            def full_invoice(result2):
                
                nonlocal stored_latest, clean
                print("stored_latest3:",stored_latest)
                
                try:
                    if type(result2['check'])==list:
                        
                        
                        for var in result2['check']:
                            all_stored = []

                            text = str(var).split("\n")

                            key = text[0].split()[0]
                            key = int(key)
                        
                            second = text[1][text[1].find(  'N' )  : ].split()
                            for i in range(len(second)):
                                second[i]=clean(second[i])
                            all_stored.extend( second[1:] )
                            forth = text[4].replace('499 - Merchandise Processing Fee','').split()
                            for i in range(len(forth)):
                                forth[i] = clean(forth[i])
                            all_stored.extend(forth)
                            fifth  = text[5].replace('501 - Harbor Maintenance Fee','').split()
                            for i in range(len(fifth)):
                                fifth[i] = clean(fifth[i])
                            all_stored.extend(fifth)
                            print("fifth5:",all_stored)
                            all_stored=[x for x in all_stored if x!='']
                            sixth = text[-1].replace('USD','').split()
                            for i in range(len(sixth)):
                                sixth[i] = clean(  sixth[i] ) 
                            all_stored.extend( sixth)


                            all_stored[7] = int(all_stored[7])
                        
                            stored_latest[str(key)] = all_stored
                        # print("stored_latest:",stored_latest)
                    elif type(result2['check'])==str:
                        all_stored = []
                        text = str(result2['check']).split("\n")
                        key = text[0].split()[0]
                        key = int(key)
                        second = text[1][text[1].find(  'N' )  : ].split()
                        for i in range(len(second)):
                            second[i]=clean(second[i])
                        
                        all_stored.extend( second[1:] )
                        forth = text[4].replace('499 - Merchandise Processing Fee','').split()
                        for i in range(len(forth)):
                            forth[i] = clean(forth[i])
                        all_stored.extend(forth)
                        fifth  = text[5].replace('501 - Harbor Maintenance Fee','').split()
                        for i in range(len(fifth)):
                            fifth[i] = clean(fifth[i])
                        all_stored.extend(fifth)
                        print("fifth6:",all_stored)
                        all_stored=[x for x in all_stored if x!='']
                        # print("fifth2:",all_stored)

                        sixth = text[-1].replace('USD','').split()
                        for i in range(len(sixth)):
                            sixth[i] = clean(  sixth[i] ) 
                        all_stored.extend( sixth)


                        all_stored[7] = int(all_stored[7])
                        # if '2323200853' in var:
                        #     print(all_stored)
                        stored_latest[str(key)] = all_stored
                    # print("stored_latest1:",stored_latest)

                except KeyError :
                        
                    
                        if type(result2['check2'])==list:
            
                            for var in result2['check2']:
                                all_stored = []
                                print("all_stored:",all_stored)

                                text = str(var).split("\n")

                                key = text[0].split()[0]
                                key = int(key)
                            
                                second = text[0][text[0].find(  'N' )  : ].split()
                                for i in range(len(second)):
                                    second[i]=clean(second[i])
                                all_stored.extend( second[1:] )
                                forth = text[3].replace('499 - Merchandise Processing Fee','').split()
                                
                                for i in range(len(forth)):
                                    forth[i] = clean(forth[i])
                                all_stored.extend(forth)
                                fifth  = text[4].replace('501 - Harbor Maintenance Fee','').split()
                                
                                for i in range(len(fifth)):
                                    fifth[i] = clean(fifth[i])
                                
                                all_stored.extend(fifth)
                                all_stored=[x for x in all_stored if x!='']
                                print("fifth7:",all_stored)
                                sixth = text[-1].replace('USD','').split()
                                for i in range(len(sixth)):
                                    sixth[i] = clean(  sixth[i] ) 
                                all_stored.extend( sixth)
                                
                                all_stored_copy = []
                                for i in range(len(all_stored)):
                                    if (all_stored[i])!='':
                                        all_stored_copy.append(  all_stored[i]  ) 

                                all_stored_copy[7] = int(all_stored_copy[7])
                                stored_latest[str(key)] = all_stored_copy
                        elif type(result2['check2'])==str:
                            all_stored = []
                            text = (result2['check2']).split("\n")
                            key = text[0].split()[0]
                            key = int(key)
                        
                            second = text[0][text[0].find(  'N' )  : ].split()
                            for i in range(len(second)):
                                second[i]=clean(second[i])
                            
                            all_stored.extend( second[1:] )
                            forth = text[3].replace('499 - Merchandise Processing Fee','').split()
                            for i in range(len(forth)):
                                forth[i] = clean(forth[i])
                            all_stored.extend(forth)
                            fifth  = text[4].replace('501 - Harbor Maintenance Fee','').split()
                            for i in range(len(fifth)):
                                fifth[i] = clean(fifth[i])
                            all_stored.extend(fifth)
                            all_stored=[x for x in all_stored if x!='']
                            print("fifth8:",all_stored)

                            sixth = text[-1].replace('USD','').split()
                            for i in range(len(sixth)):
                                sixth[i] = clean(  sixth[i] ) 
                            all_stored.extend( sixth)


                            all_stored_copy = []
                            for i in range(len(all_stored)):
                                if (all_stored[i])!='':
                                    all_stored_copy.append(  all_stored[i]  ) 

                            all_stored_copy[7] = int(all_stored_copy[7])
                            stored_latest[str(key)] = all_stored_copy



            checking_stored ={}
            def missing(result2):
                    try:
                        nonlocal checking_stored 
                        if type(result2['missingboxes'])==list:
                            for var in result2['missingboxes']:
                                all_stored = []

                                text = str(var).split("\n")

                                key = text[0].split()[0]
                                key = int(key)
                            
                                second = text[1][text[1].find(  'N' )  : ].split()
                                for i in range(len(second)):
                                    second[i]=clean(second[i])
                                all_stored.extend( second[1:] )
                                forth = text[4].replace('499 - Merchandise Processing Fee','').split()
                                for i in range(len(forth)):
                                    forth[i] = clean(forth[i])
                                all_stored.extend(forth)
                                fifth  = text[5].replace('501 - Harbor Maintenance Fee','').split()
                                for i in range(len(fifth)):
                                    fifth[i] = clean(fifth[i])
                                all_stored.extend(fifth)
                                all_stored=[x for x in all_stored if x!='']
                                print("fifth9:",all_stored)

                                checking_stored[str(key)] = all_stored
                        elif  type(result2['missingboxes'])==str:
                            all_stored = []
                            var  = result2['missingboxes']

                            text = str(var).split("\n")

                            key = text[0].split()[0]
                            key = int(key)
                            
                            second = text[1][text[1].find(  'N' )  : ].split()
                            for i in range(len(second)):
                                second[i]=clean(second[i])
                            all_stored.extend( second[1:] )
                            forth = text[4].replace('499 - Merchandise Processing Fee','').split()
                            for i in range(len(forth)):
                                forth[i] = clean(forth[i])
                            all_stored.extend(forth)
                            fifth  = text[5].replace('501 - Harbor Maintenance Fee','').split()
                            for i in range(len(fifth)):
                                fifth[i] = clean(fifth[i])
                            all_stored.extend(fifth)
                            print("fifth10:",all_stored)
                            all_stored=[x for x in all_stored if x!='']
                            # print("All_Stored:",all_stored)

                            checking_stored[str(key)] = all_stored
                    except KeyError: 
                        try:
                            stored_update = {}
                            if type(result2['missingboxes2'])==list:
                                for var in result2['missingboxes2']:
                                    all_stored = []

                                    text = str(var).split("\n")

                                    key = text[0].split()[0]
                                    key = int(key)
                                
                                    second = text[0][text[0].find(  'N' )  : ].split()
                                    for i in range(len(second)):
                                        second[i]=clean(second[i])
                                    all_stored.extend( second[1:] )
                                    forth = text[3].replace('499 - Merchandise Processing Fee','').split()
                                    for i in range(len(forth)):
                                        forth[i] = clean(forth[i])
                                    all_stored.extend(forth)
                                    fifth  = text[4].replace('501 - Harbor Maintenance Fee','').split()
                                    for i in range(len(fifth)):
                                        fifth[i] = clean(fifth[i])
                                    all_stored.extend(fifth)
                                    all_stored=[x for x in all_stored if x!='']
                                    print("fifth11:",all_stored)

                                stored_update[str(key)] = all_stored
                            elif  type(result2['missingboxes2'])==str:
                                all_stored = []
                                var  = result2['missingboxes2']


                                text = str(var).split("\n")

                                key = text[0].split()[0]
                                key = int(key)
                                
                                second = text[0][text[0].find(  'N' )  : ].split()
                                for i in range(len(second)):
                                    second[i]=clean(second[i])
                                all_stored.extend( second[1:] )
                                forth = text[3].replace('499 - Merchandise Processing Fee','').split()
                                for i in range(len(forth)):
                                    forth[i] = clean(forth[i])
                                all_stored.extend(forth)
                                fifth  = text[4].replace('501 - Harbor Maintenance Fee','').split()
                                for i in range(len(fifth)):
                                    fifth[i] = clean(fifth[i])
                                all_stored.extend(fifth)
                                all_stored=[x for x in all_stored if x!='']
                                print("fifth12:",all_stored)
                                stored_update[str(key)] = all_stored
                            checking_stored.update(stored_update)
                        except :
                            pass
            
                       
            paired_stored={}
            try:             
                (handle_sep(result2))
                print("ssssss:",stored)      

                        
                full_invoice(result2)
                # print("Full_Invoice:",Full_Invoice)

                missing(result2)
                # print("missed:",missed)

                for key , value in checking_stored.items():
                    print("Checking_Stored:",checking_stored)
                    
                    if key not in list(stored_latest.keys()):
                        print("stored_latest_key:",stored_latest)

                        
                        stored_latest[key] = value
                # print("stored_latest:",stored_latest)
                        

                stored_latest = dict(sorted(stored_latest.items(), key=lambda x: int(x[0])))

                print("stored_latest1:",stored_latest)
                # stored_latest = {k: [0.0 if x == 'FREE' else x for x in v] for k, v in stored_latest.items()}
                stored_latest = {k: [0.0 if x == 'FREE' else x for x in v if x != ''] for k, v in stored_latest.items()}

                print("stored_latest2:",stored_latest)

                
                list_which_arepaired = []
                pair  = 1 
                for key, value in  stored_latest.items():
                    if len(value)==7:
                        dict_value = {  key : value ,  str(int(key)+1): stored_latest[str(int(key)+1)]  }
                        print("dict_value:",dict_value)
                        list_which_arepaired.append(str(stored_latest[str(int(key)+1)][7]))
                        paired_stored[str(pair)] = dict_value
                        pair+=1 

                    
                

                paired_stored = dict(sorted(paired_stored.items(), key=lambda x: int(x[0])))
                # print("paired_stored:",paired_stored)
               
            except Exception as e :
                print("Error:",e)
                writeLog("Error in new template:"+str(e))



            ## these variables  are for totaling
            MPF_total = float() ; HMF_total=float();Total_Entered_Value=float(); Total_Duty = float()
            print("MPF_total4:",MPF_total)
            print("Total_Duty1:",Total_Duty)

            # print("Total_Entered_Value:",Total_Entered_Value)

            try:
                ## collecting data of Comm invoice number and FOB values 
                ## and storing in stored_info_txt variable 
                #stored_info_txt variable is stored data of invoice number and FOB Values
                if not os.path.exists("comm_invoice_fob.txt"):
                    text_box.insert(tk.END, "Invoice numbers and FOB values are null" "\n")


                with open("comm_invoice_fob.txt","r") as file:
                    stored_info_txt = {}
                    for value  in file:
                        content = value.split()

                        stored_info_txt[(content [0])] = float(content [1].replace("\n","").replace("\r","").replace(',','').replace('-','') )
            except Exception as e:
                text_box.insert(tk.END,str(e) + "\n")

                return
            #------------
            
            try:
                stored = {}
                ## copying work 
                for key,value in stored_latest.items():
                    # print("keys:",key)
                    print("dict_of_ALL:",key,value)
                    if len(value)>7:
                        stored[str(int(value[7]))] = value
                print("stored2:",stored)
                    
                
                #----------
            except  Exception as e:
                writeLog('error while copying data '+str(e))

            try:
                ## fetching both available and not available invoices 
                list_comm_invoice = list(stored_info_txt.keys())
                print("list_comm_invoice:",list_comm_invoice)
                available_invoices=[]
                not_available_invoices = []
                for key, value   in stored.items():
                    # print("Stored_Items:",stored)
                    if key in list_comm_invoice:
                       
                        available_invoices.append(key)
                        print("available_invoices:",available_invoices)


                    else:
                        not_available_invoices.append(key)
                        print("not_available_invoices:",not_available_invoices)

                #-----
            except  Exception as e:
                writeLog('Error while fetching both available and not available invoices '+str(e))
            
            problem_count  =  0 

            
            with open("reasons.txt",'w') as file : 
                ## if any invoice not available then message will show on textbox
                for current in not_available_invoices:
                    text_box.insert(tk.END, f"{int(current)} this invoice not available in your PDF FOLDER" + "\n")
                    problem_count+= 1 
                    file.write(f"{int(current)} this invoice not available in your PDF FOLDER" + "\n")

                for value  in available_invoices:
                    # print("value:",value)
                    
                   
                    try:
                        FOB = int((float(stored_info_txt[value])))
                        # print("fob1:",FOB)
                        
                        if value not in list_which_arepaired:
                            issue = False
                    
                            forprint = [value ,[]]
                            # print("forprint:",forprint)
                            if len(stored[value])==12:
                                list_value = stored[value]
                                Invoice_Value  = list_value[-4]
                                mmv = list_value[-3]
                                last_entered_value = list_value[-1]
                                if int(last_entered_value) !=FOB :

                                    if int(last_entered_value )  ==  FOB+1 :
                                        pass
                                    else:
                                        issue = True
                                        forprint[1].append('Entered Value not matched1')
                                        print(  )
                                    
                                if int(Invoice_Value+mmv)!= int(last_entered_value):
                                    issue = True
                                    forprint[1].append(f"After MMV applied  {Invoice_Value} with {mmv} the entered value is not matched {int(last_entered_value)}")
                                entered_at32 = list_value[0]
                                if int(entered_at32)!=int(last_entered_value):
                                    issue = True
                                    forprint[1].append('Entered Value not matched which is at 32 A.Entered Value  ')
                                rate1= list_value[1]
                                ratecal= list_value[2]
                                mpfrate = list_value[3]
                                mpfcal = list_value[4]
                                hmfrate = list_value[5]
                                hmfcal = list_value[6]
                                MPF_total+= mpfcal
                                print("MPF_total3:",MPF_total)

                                HMF_total+=hmfcal
                                Total_Duty+= ratecal
                                print("Total_Duty2:",Total_Duty)
                                Total_Entered_Value+=last_entered_value
                                print("total entered value:",Total_Entered_Value)
                                pythonratecal = float(entered_at32*rate1/100)   ;pythonmpfcal=float(entered_at32*mpfrate/100);pythonhmfcal=float(entered_at32*hmfrate/100)

                                if int(pythonratecal)!=int(ratecal):
                                    if int(pythonratecal)+1==int(ratecal):pass
                                    else:
                                        issue = True
                                        msg = f'This calculated value1 {ratecal}  not matched with  Invoice value {entered_at32} % rate {rate1}'
                                        forprint[1].append(msg)
                                
                                if int(pythonmpfcal)!=int(mpfcal):
                                    if int(pythonmpfcal)+1==int(mpfcal):pass
                                    else:
                                        issue = True
                                        msg = f'This calculated value2 {mpfcal}  not matched with  Invoice value {entered_at32} % rate {mpfrate}'
                                        forprint[1].append(msg)
                                
                                if int(pythonhmfcal)!=int(hmfcal):
                                    if int(pythonhmfcal)+1==int(hmfcal):pass
                                    else:
                                        issue = True
                                        msg = f'This calculated value3 {hmfcal}  not matched with  Invoice value {entered_at32} % rate {hmfrate}'
                                        forprint[1].append(msg)
                            elif len(stored[value])==11:
                                list_value = stored[value]
                                print("list_value:",list_value)
                                Invoice_Value  = list_value[-3]
                                
                                last_entered_value = list_value[-1]
                                # print("last_entered_value:",last_entered_value)
                                # print("FOB:",FOB)
                                if int(last_entered_value) !=FOB:
                                    if int(last_entered_value)  ==FOB+1:
                                        
                                        pass
                                    else:

                                        issue = True
                                        forprint[1].append('Entered Value not matched2')
                                
                                entered_at32 = list_value[0]
                                print("entered_at32:",entered_at32)
                                if int(entered_at32)!=int(last_entered_value):
                                    issue = True
                                    forprint[1].append('Entered Value not matched which is at 32 A.Entered Value  ')
                                rate1= list_value[1]
                                ratecal= list_value[2]
                                mpfrate = list_value[3]
                                mpfcal = list_value[4]

                                hmfrate = list_value[5]
                                print("hmfrate:",hmfrate)
                                hmfcal = list_value[6]
                                MPF_total+= mpfcal
                                print("MPF_total2:",MPF_total)
                                HMF_total+=hmfcal
                                Total_Duty+= ratecal
                                print("Total_Duty3:",Total_Duty)
                                Total_Entered_Value+=last_entered_value
                                pythonratecal = float(entered_at32*rate1/100)   ;pythonmpfcal=float(entered_at32*mpfrate/100);pythonhmfcal=float(entered_at32*hmfrate/100)

                                if int(pythonratecal)!=int(ratecal):
                                    if int(pythonratecal)+1==int(ratecal):pass
                                    else:
                                        issue = True
                                        msg = f'This calculated value4 {ratecal}  not matched with  Invoice value {entered_at32} % rate {rate1}'
                                        forprint[1].append(msg)
                                
                                if int(pythonmpfcal)!=int(mpfcal):
                                    if int(pythonmpfcal)+1==int(mpfcal):pass
                                    else:
                                        issue = True
                                        msg = f'This calculated value5 {mpfcal}  not matched with  Invoice value {entered_at32} % rate {mpfrate}'
                                        forprint[1].append(msg)
                                
                                if int(pythonhmfcal)!=int(hmfcal):
                                    print("pythonhmfcal:",pythonhmfcal)
                                    print("hmfcal:",hmfcal)
                                    if int(pythonhmfcal)+1==int(hmfcal):pass
                                    else:
                                        issue = True
                                        msg = f'This calculated value6 {hmfcal}  not matched with  Invoice value {entered_at32} % rate {hmfrate}'
                                        forprint[1].append(msg)



                        ## if any issue raise of  any invoice then message will show on textbox
                            if not issue:
                                text_box.insert(tk.END, f"{value}   OK" +"\n")
                            if issue:
                                problem_count += 1 
                                text_box.insert(tk.END, f"{forprint[0]}   NOT OK" +"\n")
                                text_box.insert(tk.END, f"Following Reasons: " +"\n")


                                for j in forprint[-1]:
                                    text_box.insert(tk.END,  f"{forprint[0]} :  {j}\n" )
                                    file.write( f"{forprint[0]} :  {j}\n"  )  
                                    problem_count+=1 
                    except Exception as e :
                        writeLog('Error in first checking '+str(e))
                    #----
                #------------------------------------------------------------------------------

                ## which invoice separeted with two pages 
                try:
                
                    for key, value in paired_stored.items():
                        
                        dict1 = value
                        first_part= []
                        second_part=[]
                        first_done =1 
                        second_done = 0 
                        for key , value in dict1.items():
                            if first_done:
                                first_part.extend(value)
                                first_done= 0 
                                second_done = 1 
                            elif second_done:
                                second_part.extend(value)
                                first_done = 0
                                second_done = 0
                        if len(second_part)==12:
                            issue =  0 
                            list_value = second_part
                            forprint = [  str(list_value[7] ) ,[] ]
                                
                            Invoice_Value  = list_value[-4]
                            mmv = list_value[-3]
                            last_entered_value = list_value[-1]
                            FOB = int(float(stored_info_txt[str(list_value[7])]))

                            if int(last_entered_value)!=FOB:
                                if int(last_entered_value)==FOB+1:
                                    pass
                                else:
                                    issue = True
                                    forprint[1].append('Entered Value not matched3')
                            if int(Invoice_Value+mmv)!= int(last_entered_value):
                                issue = True
                                forprint[1].append(f"After MMV applied  {Invoice_Value} with {mmv} the entered value is not matched {int(last_entered_value)}")
                            var1= first_part[0]
                            var2 = second_part[0]
                            if int(var1+var2)!=int(last_entered_value):
                                issue=True
                                forprint[1].append(f'Sum of  {var1} and {var2} is not matched with entered value {last_entered_value}')
                            enteredvaue32_first = first_part[0]
                            rat1first = first_part[1]
                            rate1calfirst = first_part[2]
                            mpfratefirst = first_part[3]
                            mpfcalfirst = first_part[4]
                            hmfratefirst = first_part[5]
                            hmfcalfirst = first_part[6]

                            enteredvaue32_second = second_part[0]
                            rat1second = second_part[1]
                            rate1calsecond = second_part[2]
                            mpfratesecond = second_part[3]
                            mpfcalsecond = second_part[4]
                            hmfratesecond = second_part[5]
                            hmfcalsecond = second_part[6]

                            pythonrate1calfirst= float(enteredvaue32_first*rat1first/100)
                            pythonmpfcalfirst= float(enteredvaue32_first*mpfratefirst/100)
                            pythonhmfcalfirst = float(enteredvaue32_first*hmfratefirst/100)

                            pythonrate1calsecond= float(enteredvaue32_second*rat1second/100)
                            pythonmpfcalsecond= float(enteredvaue32_second*mpfratesecond/100)
                            pythonhmfcalsecond = float(enteredvaue32_second*hmfratesecond/100)

                            MPF_total+= (mpfcalfirst+ mpfcalsecond  )
                            print("MPF_total1:",MPF_total)
                            HMF_total+= (hmfcalfirst+hmfcalsecond)
                            Total_Duty+=( rate1calfirst+rate1calsecond)
                            print("Total_Duty4:",Total_Duty)
                            Total_Entered_Value+= last_entered_value

                            # writeLog(f'totaling ' +str(Total_Duty) +"---"  +str(rate1calfirst)  +"----"  +str(rat1second))

                            if int(pythonrate1calfirst)!=int(rate1calfirst):
                                if int(pythonrate1calfirst)+1==int(rate1calfirst):
                                    pass
                                else:
                                    issue = True
                                    msg = f'This calculated value7 {rate1calfirst}  not matched with  Invoice value {enteredvaue32_first} % rate {rat1first}'
                                    forprint[1].append(msg)
                                    pass
                            if int(pythonmpfcalfirst)!=int(mpfcalfirst):
                                if int(pythonmpfcalfirst)+1==int(mpfcalfirst):
                                    pass
                                else:
                                    issue = True
                                    msg = f'This calculated value8 {mpfcalfirst}  not matched with  Invoice value {enteredvaue32_first} % rate {mpfratefirst}'
                                    forprint[1].append(msg)
                            if int(pythonhmfcalfirst)!=int(hmfcalfirst):
                                if int(pythonhmfcalfirst)+1==int(hmfcalfirst):
                                    pass
                                else:
                                    issue = True
                                    msg = f'This calculated value10 {hmfcalfirst}  not matched with  Invoice value {enteredvaue32_first} % rate {hmfratefirst}'
                                    forprint[1].append(msg)


                            if int(pythonrate1calsecond)!=int(rate1calsecond):
                                if int(pythonrate1calsecond)+1==int(rate1calsecond):
                                    pass
                                else:
                                    issue = True
                                    msg = f'This calculated value11 {rate1calsecond}  not matched with  Invoice value {enteredvaue32_second} % rate {rat1second}'
                                    forprint[1].append(msg)
                            
                            if int(pythonmpfcalsecond)!=int(mpfcalsecond):
                                if int(pythonmpfcalsecond)+1==int(mpfcalsecond):
                                    pass
                                else:
                                    issue = True
                                    msg = f'This calculated value12 {mpfcalsecond}  not matched with  Invoice value {enteredvaue32_second} % rate {mpfratesecond}'
                                    forprint[1].append(msg)
                            if int(pythonhmfcalsecond)!=int(hmfcalsecond):
                                if int(pythonhmfcalsecond)+1==int(hmfcalsecond):
                                
                                    pass
                                else:
                                    issue = True
                                    msg = f'This calculated value13 {hmfcalsecond}  not matched with  Invoice value {enteredvaue32_second} % rate {hmfratesecond}'
                                    forprint[1].append(msg)
                            

                            if not issue:
                                text_box.insert(tk.END, f"{str(list_value[7] )}   OK" +"\n")
                            if issue:
                                problem_count += 1 
                                text_box.insert(tk.END, f"{str(list_value[7] )}   NOT OK" +"\n")
                                text_box.insert(tk.END, f"Following Reasons: " +"\n")


                                for j in forprint[-1]:
                                    text_box.insert(tk.END,  f"{str(list_value[7] )} :  {j}\n" )
                                    file.write( f"{str(list_value[7] )} :  {j}\n"  )  
                                    problem_count+=1
                        elif len(second_part)==11:
                            issue =  0 
                            list_value = second_part
                            forprint = [  str(list_value[7] ) ,[] ]

                                
                            Invoice_Value  = list_value[8]
                            
                            last_entered_value = list_value[-1]
                            FOB = int((float(stored_info_txt[str(list_value[7])])))

                            if int(last_entered_value)!=FOB:
                                if int(last_entered_value)==FOB+1:
                                    pass
                                else:
                                    issue = True
                                    forprint[1].append('Entered Value not matched4')
                            
                            var1= first_part[0]
                            var2 = second_part[0]

                            if int(var1+var2)!=int(last_entered_value):
                                issue=True
                                forprint[1].append(f'Sum of  {var1} and {var2} is not matched with entered value {last_entered_value}')
                            enteredvaue32_first = first_part[0]
                            rat1first = first_part[1]
                            rate1calfirst = first_part[2]
                            mpfratefirst = first_part[3]
                            mpfcalfirst = first_part[4]
                            hmfratefirst = first_part[5]
                            hmfcalfirst = first_part[6]

                            enteredvaue32_second = second_part[0]
                            rat1second = second_part[1]
                            rate1calsecond = second_part[2]
                            mpfratesecond = second_part[3]
                            mpfcalsecond = second_part[4]
                            hmfratesecond = second_part[5]
                            hmfcalsecond = second_part[6]

                            pythonrate1calfirst= float(enteredvaue32_first*rat1first/100)
                            pythonmpfcalfirst= float(enteredvaue32_first*mpfratefirst/100)
                            pythonhmfcalfirst = float(enteredvaue32_first*hmfratefirst/100)

                            pythonrate1calsecond= float(enteredvaue32_second*rat1second/100)
                            pythonmpfcalsecond= float(enteredvaue32_second*mpfratesecond/100)
                            pythonhmfcalsecond = float(enteredvaue32_second*hmfratesecond/100)

                            MPF_total+= (mpfcalfirst+ mpfcalsecond  )
                            print("MPF_total1:",MPF_total)
                            HMF_total+= (hmfcalfirst+hmfcalsecond)
                            Total_Duty+=( rate1calfirst+rate1calsecond)
                            Total_Entered_Value+= last_entered_value

                            # writeLog(f'totaling ' +str(Total_Duty) +"---"  +str(rate1calfirst)  +"----"  +str(rat1second))

                            if int(pythonrate1calfirst)!=int(rate1calfirst):
                                if int(pythonrate1calfirst)+1==int(rate1calfirst):
                                    pass
                                else:
                                    issue = True
                                    msg = f'This calculated value13 {rate1calfirst}  not matched with  Invoice value {enteredvaue32_first} % rate {rat1first}'
                                    forprint[1].append(msg)
                                    pass
                            if int(pythonmpfcalfirst)!=int(mpfcalfirst):
                                if int(pythonmpfcalfirst)+1==int(mpfcalfirst):
                                    pass
                                else:
                                    issue = True
                                    msg = f'This calculated value14 {mpfcalfirst}  not matched with  Invoice value {enteredvaue32_first} % rate {mpfratefirst}'
                                    forprint[1].append(msg)
                            if int(pythonhmfcalfirst)!=int(hmfcalfirst):
                                if int(pythonhmfcalfirst)+1==int(hmfcalfirst):
                                    pass
                                else:
                                    issue = True
                                    msg = f'This calculated value15 {hmfcalfirst}  not matched with  Invoice value {enteredvaue32_first} % rate {hmfratefirst}'
                                    forprint[1].append(msg)


                            if int(pythonrate1calsecond)!=int(rate1calsecond):
                                if int(pythonrate1calsecond)+1==int(rate1calsecond):
                                    pass
                                else:
                                    issue = True
                                    msg = f'This calculated value16 {rate1calsecond}  not matched with  Invoice value {enteredvaue32_second} % rate {rat1second}'
                                    forprint[1].append(msg)
                            
                            if int(pythonmpfcalsecond)!=int(mpfcalsecond):
                                if int(pythonmpfcalsecond)+1==int(mpfcalsecond):
                                    pass
                                else:
                                    issue = True
                                    msg = f'This calculated value17 {mpfcalsecond}  not matched with  Invoice value {enteredvaue32_second} % rate {mpfratesecond}'
                                    forprint[1].append(msg)
                            if int(pythonhmfcalsecond)!=int(hmfcalsecond):
                                if int(pythonhmfcalsecond)+1==int(hmfcalsecond):
                                
                                    pass
                                else:
                                    issue = True
                                    msg = f'This calculated value18 {hmfcalsecond}  not matched with  Invoice value {enteredvaue32_second} % rate {hmfratesecond}'
                                    forprint[1].append(msg)
                            

                            if not issue:
                                text_box.insert(tk.END, f"{str(list_value[7] )}   OK" +"\n")
                            if issue:
                                problem_count += 1 
                                text_box.insert(tk.END, f"{str(list_value[7] )}   NOT OK" +"\n")
                                text_box.insert(tk.END, f"Following Reasons: " +"\n")


                                for j in forprint[-1]:
                                    text_box.insert(tk.END,  f"{str(list_value[7] )} :  {j}\n" )
                                    file.write( f"{str(list_value[7] )} :  {j}\n"  )  
                                    problem_count+=1
                        
                except Exception as e :
                    writeLog('Error in checking in second '+str(e))
               
                try:
                    print("MPF_Total_:",MPF_total)
                    pdf_total_MPF=int(result['MPF'])
                    if MPF_total<31.67:
                        MPF_total=31.67
                except Exception as e:
                    writeLog('Error in MPF working '+str(e))
            
                try:
                    print("MPF_total:",MPF_total)
                    ## Total MPF checking 
                    if pdf_total_MPF!= int(MPF_total):
                        print("pdf_total_MPF:",pdf_total_MPF)
                        print("MPF_total:",MPF_total)
                        # print("pdf_total_MPF:",pdf_total_MPF)
                        text_box.insert(tk.END, f"Total MPF not matched1 " +"\n")
                        problem_count += 1
                        file.write(f"Total MPF not matched2" +"\n")
                    #----
                except Exception as e:
                    writeLog('Error while Total MPF checking '+str(e))

                try:
                    ## Total HMF checking
                    if int(result['HMF'])!= int(HMF_total):
                        print("HMF_TOTAL:",HMF_total)
                        print("HMF:",result['HMF'])
                        text_box.insert(tk.END, f"Total HMF not matched1" +"\n")
                        problem_count += 1
                        file.write(f"Total HMF not matched2" +"\n")
                    #----
                except Exception as e:
                    writeLog('Error while Total HMF checking '+str(e))
                
                try:
                    ## Total Total Entered Value  checking
                    if int(result['Total Entered Value'])!= int(Total_Entered_Value):
                        print("totalvalue:",result['Total Entered Value'])
                        print("Total_Entered_Value:",Total_Entered_Value)
                        text_box.insert(tk.END, f"Total Entered Value not matched1" +"\n")
                        problem_count += 1
                        file.write(f"Total Entered Value not matched2" +"\n" )
                    #-----
                except Exception as e:
                    writeLog('Error while Total Entered Value checking '+str(e))
                
                try:
                    ## Total Other Fees checking
                    if int((MPF_total+HMF_total  ))!= int(result['Total Other Fees']):
                        text_box.insert(tk.END, f"Total Other Fees not matched" +"\n")
                        problem_count += 1
                        file.write(f"Total Other Fees not matched" +"\n")
                    #-----
                except Exception as e:
                    writeLog('Error while Total Other Fees checking '+str(e))

                try:
                    ## Duty checking
                    Total_Duty += sum(dollar_values) 
                    print("Result_Duty:",result['Duty'])
                    print("Total_Duty_sum:",Total_Duty)
                    if int(Total_Duty)!= int(result['Duty']):
                        # print("Result_Duty:",result['Duty'])
                        # print("Total_Duty:",Total_Duty)
                        text_box.insert(tk.END, f"Duty not matched" +"\n")
                        problem_count += 1
                        file.write(f"Duty not matched" +"\n")
                    #-----
                except Exception as e:
                    writeLog('Error while Duty checking '+str(e))
                


                try:
                    ## Other checking
                    if int((MPF_total+HMF_total))!= int(result['Other']):
                        text_box.insert(tk.END, f"Other not matched" +"\n")
                        problem_count += 1
                        file.write(f"Other not matched" +"\n" )
                    #------
                except Exception as e:
                    writeLog('Error while Other checking '+str(e))
                
                try:
                    ## Last Total checking 
                    if int((Total_Duty+(MPF_total+HMF_total)))!= int(result['Total']):
                        text_box.insert(tk.END, f"Total not matched" +"\n")
                        problem_count += 1
                        file.write(f"Last Total not matched" +"\n")
                    #-----
                except Exception as e:
                    writeLog('Error while Last total checking '+str(e))
                
                ## if any problem raise then it will shows all issues and give reason.txt 
                if problem_count==0:
                    text_box.insert(tk.END, f"Overall OK" +"\n")
                    text_box.insert(tk.END, f"    " +"\n")

                elif problem_count>= 1:
                    text_box.insert(tk.END, f"Overall NOT OK " +"\n")
                    text_box.insert(tk.END, f"    " +"\n")

                    def download_file():
                        file_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text files", "*.txt")])
                        if file_path:
                            with open("reasons.txt", "r") as source_file:
                                with open(file_path, "w") as target_file:
                                    target_file.write(source_file.read())
                    if problem_count!=0:
                        download_button = tk.Button(console_frame  ,  text="Download reason.txt", command=download_file)
                        download_button.grid(row=3 , column=0)
                #--------
                            
            
                
        ##GUI of validate amount 
        root = tk.Tk()
        root.title("Validate Invoice(SIPL)")
        root.iconbitmap("content/SIPL.ico")  # Ensure the path to your icon file is correct
        root.geometry('520x460')

        panel_color = '#E0E0E0'
        button_font = ("Arial", 10, "bold")

        # Console file section
        console_frame = tk.Frame(root)
        console_frame.pack(pady=10)

        console_label = tk.Label(console_frame, text="Console File:")
        console_label.grid(row=0, column=0, sticky=tk.W)

        console_entry = tk.Entry(console_frame, width=50)
        console_entry.grid(row=0, column=1, sticky=tk.W)

        console_browse_button = tk.Button(console_frame, text="Browse", command=lambda: browse_file(console_entry))
        console_browse_button.grid(row=0, column=3, padx=5, sticky=tk.W)

        # PDF folder section
        pdf_label = tk.Label(console_frame, text="PDF Folder:")
        pdf_label.grid(row=2, column=0, sticky=tk.W, pady=5)

        pdf_entry = tk.Entry(console_frame, width=50)
        pdf_entry.grid(row=2, column=1, sticky=tk.W, pady=5)

        pdf_browse_button = tk.Button(console_frame, text="Browse", command=lambda: Adding_invoice_and_FOB(pdf_entry))
        pdf_browse_button.grid(row=2, column=3, padx=5, sticky=tk.W, pady=5)

        process_button = tk.Button(console_frame, text="Start Process", command=StartProcess, bg="#fa8f16", fg="#fff")
        process_button.grid(row=3, column=1, padx=5, sticky=tk.W, pady=5)

        # Text box with scrollbar
        text_frame = tk.Frame(root)
        text_frame.pack(pady=10)

        text_box = tk.Text(text_frame, height=19, width=78, font=("Helvetica", 8))
        text_box.grid(row=3, column=0, columnspan=3)

        scrollbar = tk.Scrollbar(text_frame, orient="vertical", command=text_box.yview)
        scrollbar.grid(row=3, column=3, sticky="ns")
        text_box.config(yscrollcommand=scrollbar.set)

        developer_label = tk.Label(root, text="Developed By: Sagarinfotech.com", font=('Helvetica', 8, 'underline'), bg=panel_color, fg='blue', cursor='hand2')
        developer_label.pack(side="left", fill="x")
        developer_label.bind("<Button-1>", on_link_click)

        root.mainloop()
        #-----------------------

# Create main window
if __name__=="__main__":
    ## if comm_invoice_fob.txt file exist then remove 
    
    #------
    
    def on_link_click(event):
        webbrowser.open("https://www.sagarinfotech.com")

    #error,_=Licence()
     
    def showMessage():
        _,message=Licence()
        lbl_Message = Label(root , text = message,bg=panel_color,fg = "red")
        lbl_Message.grid(row=0, column=0, sticky="w", padx=15, pady=5)

    # create root window
    root = Tk()
    root.title("Welcome to AI World(SIPL)")
    root.iconbitmap("content/SIPL.ico")
    # Set geometry (widthxheight) 
    root.geometry('450x185') 

    panel_color = '#E0E0E0'

    # panel = Frame(root, bg=panel_color, padx=5, pady=5)
    # panel.grid(row=0, column=0, padx=5, pady=5, sticky='nsew')

    #adding a label to the root window
  
    root.after(5000, showMessage)
    textff= tk.Label( root  , text=" ")
    textff.grid ( row = 1 , column= 0)

    # Create a button with specified font properties
    button_font = ("Arial", 10, "bold")
    btn = Button( root , text = "REDUCE INVOICE ",bg="lightblue", fg="black", padx=10, pady=5, relief="raised", font=("Helvetica", 12, "bold"), command=reduce_invoice)
    btn.grid(row=2, column=0, sticky="w", padx=15, pady=5)
    btn1 = Button(root  , text = "VALIDATE AMOUNT ",bg="lightgreen", fg="black", padx=10, pady=5, relief="raised", font=("Helvetica", 12, "bold"), command=validate_amount)
    btn1.grid(row=2, column=1, sticky="w", padx=15, pady=5)
    root.resizable(width=False, height=False)
    text1 = tk.Label(root , text = " ")
    text1.grid(row=3,column= 0 )
    text2 = tk.Label(root , text = " ")
    text2.grid(row=4,column= 0 )
    text3 = tk.Label(root , text = " ")
    text3.grid(row=5,column= 0 )

    label_lnk = Label(root, text="Developed By: Sagarinfotech.com", font=('Helvetica', 8, 'underline'), bg=panel_color, fg='blue', cursor='hand2')
    label_lnk.grid(row=6, column=0, padx=1, pady=1, columnspan=2, sticky='sw')
    label_lnk.bind("<Button-1>", on_link_click)

    # Execute Tkinter
    root.mainloop()


#   pyinstaller --onefile --icon=content/SIPL.ico  pdf_converter.py
#   pyinstaller --onefile --icon=content/SIPL.ico --noconsole  pdf_converter.py